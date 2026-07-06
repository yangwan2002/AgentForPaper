"""原稿就地增补（inplace-augment-sections，方案 C）。

在用户原文件（.docx/.tex）上**插入**新章节与参考文献，**从不 re-emit 原有内容**——因此
原稿的公式（OMML / ``\\begin{equation}``）、表格、字体/样式/编号/页眉页脚、preamble/宏等
**逐字保留**。这是对 ``polish_docx_inplace`` / ``polish_latex_inplace``「只改不加」保结构范式
的扩展：从「只改」扩展到「能加新章节」，且同样**只增不改（Additive_Only）**。

本模块 Task 1 提供数据模型与 :class:`InplaceLatexAugmenter`（纯文本、易测、无 docx 依赖）；
Task 2 提供 :class:`InplaceDocxAugmenter`（python-docx 插入 + 无损校验 + 参考文献排版）。

设计红线（见 spec）：
- **只增不改**：只新增元素，绝不重写/删除原有内容。
- **失败保留原稿**：校验不过即判失败、保留原稿，绝不交付破坏性产物。
"""

from __future__ import annotations

import re
from dataclasses import dataclass, field


@dataclass
class SectionSpec:
    """待插入的新章节规格。

    - ``title``：章节标题（如「引言」）。
    - ``body``：章节正文（纯文本/Markdown；按空行分段）。
    - ``position``：``"start"`` 插到正文开头（首个 section 前）；``"end"`` 插到正文末尾
      （参考文献之前）。``anchor`` 提供时优先按锚点定位（docx）。
    - ``anchor``：可选锚文本，插到匹配段落前（docx 用）。
    """

    title: str
    body: str = ""
    position: str = "start"
    anchor: str | None = None


@dataclass
class AugmentResult:
    """一次就地增补的结果。"""

    ok: bool
    out_path: str = ""
    inserted_sections: int = 0
    inserted_references: int = 0
    notes: list[str] = field(default_factory=list)
    error: str = ""


# --------------------------------------------------------------------------- #
# LaTeX 就地增补（纯文本插入，只增不改）
# --------------------------------------------------------------------------- #

# 定位 \begin{document} 结束位置（含）；用于「无 \section 时插到正文开头」。
_BEGIN_DOC = re.compile(r"\\begin\{document\}")
# 定位 \end{document} 起始位置；用于「参考文献插到文末之前」。
_END_DOC = re.compile(r"\\end\{document\}")
# 定位首个 \section{...}（新章节默认插到它前面）。
_FIRST_SECTION = re.compile(r"\\section\*?\s*\{")
# 已存在参考文献块（避免重复插入）。
_HAS_BIB = re.compile(r"\\begin\{thebibliography\}|\\bibliography\s*\{")


class InplaceLatexAugmenter:
    """LaTeX 就地增补器：在原 .tex 源码的插入点新增章节/参考文献，其余逐字保留。"""

    def augment(
        self,
        source: str,
        *,
        sections: "list[SectionSpec] | tuple[SectionSpec, ...]" = (),
        references: "list[str] | tuple[str, ...]" = (),
    ) -> tuple[str, AugmentResult]:
        """在 ``source`` 上就地增补，返回 ``(新源码, AugmentResult)``。

        只在插入点新增文本；原文其余部分逐字保留（子串校验保证）。任一步失败 → 返回原文
        与 ``ok=False``，绝不产出破坏原文的结果。
        """
        result = AugmentResult(ok=True)
        text = source or ""
        try:
            for spec in sections or ():
                text = self._insert_section(text, spec, result)
            if references:
                text = self._append_references(text, list(references), result)
        except Exception as exc:  # noqa: BLE001 - 增补异常不抛，诚实回退原文
            return source, AugmentResult(
                ok=False, error=f"latex 增补异常：{type(exc).__name__}: {exc}"
            )

        # 子串校验（Property 5）：原文的每个非插入字符都应原位保留。宽松实现——
        # 断言原文去除所有空白后，是产物去除所有空白后的子序列前缀不可行；改用更强的
        # 直接保证：我们只做「在某偏移插入字符串」，故原文必是产物删去插入片段后的结果。
        # 这里做一次防御式校验：原文长度 <= 产物长度，且原文按行仍全部出现。
        if not self._preserved(source or "", text, result):
            return source, AugmentResult(
                ok=False,
                error="latex 增补后原文未被逐字保留，已放弃增补、保留原稿。",
                notes=result.notes,
            )
        return text, result

    def _insert_section(self, text: str, spec: SectionSpec, result: AugmentResult) -> str:
        """在首个 \\section 前（无则 \\begin{document} 后、再无则文首）插入新章节。"""
        block = self._render_section(spec)
        if spec.position == "end":
            pos = self._end_insert_pos(text)
        else:
            m = _FIRST_SECTION.search(text)
            if m is not None:
                pos = m.start()
            else:
                begin = _BEGIN_DOC.search(text)
                pos = begin.end() + 1 if begin is not None else 0
                if begin is None:
                    result.notes.append("未找到 \\begin{document}，章节插入到文首")
        new_text = text[:pos] + block + text[pos:]
        result.inserted_sections += 1
        return new_text

    def _append_references(
        self, text: str, entries: list[str], result: AugmentResult
    ) -> str:
        """在 \\end{document} 前插入唯一一份 thebibliography；已有则跳过。"""
        if _HAS_BIB.search(text):
            result.notes.append("原文已含参考文献块，未重复插入")
            return text
        block = self._render_bibliography(entries)
        pos = self._end_insert_pos(text)
        new_text = text[:pos] + block + text[pos:]
        result.inserted_references += len(entries)
        return new_text

    @staticmethod
    def _end_insert_pos(text: str) -> int:
        """返回 \\end{document} 前的插入偏移；无则文末（安全回退）。"""
        m = _END_DOC.search(text)
        return m.start() if m is not None else len(text)

    @staticmethod
    def _render_section(spec: SectionSpec) -> str:
        star = ""
        title = spec.title.strip()
        body = (spec.body or "").strip()
        return f"\n\\section{star}{{{title}}}\n{body}\n\n"

    @staticmethod
    def _render_bibliography(entries: list[str]) -> str:
        lines = ["\n\\begin{thebibliography}{99}"]
        for i, entry in enumerate(entries, start=1):
            text = (entry or "").strip().replace("\n", " ")
            lines.append(f"\\bibitem{{ref{i}}} {text}")
        lines.append("\\end{thebibliography}\n\n")
        return "\n".join(lines)

    @staticmethod
    def _preserved(original: str, produced: str, result: AugmentResult) -> bool:
        """校验原文逐字保留于产物：产物删去所有新插入片段后应 == 原文。

        实现：我们的插入都是「在偏移处插入连续片段」，等价于 ``produced`` 是把若干片段
        插入 ``original`` 得到的超序列。用「``original`` 是 ``produced`` 的子序列且长度
        关系成立」做防御式必要条件校验（充分性由插入实现本身保证）。
        """
        if len(produced) < len(original):
            return False
        # original 必须是 produced 的子序列（逐字符顺序保留）。
        it = iter(produced)
        return all(ch in it for ch in original)


# --------------------------------------------------------------------------- #
# DOCX 就地增补（python-docx 插入 + 无损校验 + 参考文献排版，只增不改）
# --------------------------------------------------------------------------- #

# 参考文献标题去重用的关键词（命中即视为已存在参考文献标题，不重复插入）。
_REF_HEADING_KW = ("参考文献", "references", "bibliography")


class InplaceDocxAugmenter:
    """DOCX 就地增补器：往原 docx 插入新章节 / 追加参考文献，原有内容逐字保留。

    只增不改：绝不重写任何既有 run；产物经 Preservation_Check（结构计数只增不减 + 原标题
    集合 ⊆ 产物标题集合）确认原稿内容无损，否则判失败、复制原稿、``ok=False``。
    """

    def augment(
        self,
        in_path: str,
        out_path: str,
        *,
        sections: "list[SectionSpec] | tuple[SectionSpec, ...]" = (),
        references: "list[str] | tuple[str, ...]" = (),
        references_heading: str = "参考文献",
    ) -> AugmentResult:
        """在原 docx 上就地增补章节/参考文献，写出到 ``out_path``（原稿只读）。"""
        try:
            import docx  # noqa: WPS433 - 可选依赖，惰性导入
        except ImportError as exc:  # pragma: no cover - 环境相关
            raise RuntimeError(
                "DOCX 就地增补需要 python-docx，请安装：pip install '.[docx]'"
            ) from exc

        from paper_agent.export.atomic_write import atomic_finalize
        from paper_agent.export.docx_structural import structural_fields

        result = AugmentResult(ok=True, out_path=out_path)
        document = docx.Document(in_path)
        pre = structural_fields(document)

        try:
            for spec in sections or ():
                self._insert_section(document, spec)
                result.inserted_sections += 1
            if references:
                inserted = self._append_references(
                    document, list(references), references_heading
                )
                result.inserted_references += inserted
        except Exception as exc:  # noqa: BLE001 - 增补异常 → 保留原稿、诚实上报
            self._copy_original(in_path, out_path)
            return AugmentResult(
                ok=False, out_path=out_path,
                error=f"docx 增补异常：{type(exc).__name__}: {exc}",
            )

        # Preservation_Check：结构计数只增不减 + 原标题集合 ⊆ 产物标题集合。
        post = structural_fields(document)
        if not self._preserved(pre, post, result):
            self._copy_original(in_path, out_path)
            return AugmentResult(
                ok=False, out_path=out_path,
                error="docx 增补破坏了原有结构，已放弃增补、保留原稿。",
                notes=result.notes,
            )

        # 原子写出到独立产物（原稿全程只读）。
        import os
        import tempfile

        directory = os.path.dirname(os.path.abspath(out_path)) or "."
        os.makedirs(directory, exist_ok=True)
        fd, tmp = tempfile.mkstemp(dir=directory, prefix=".tmp_aug_", suffix=".docx")
        os.close(fd)
        try:
            document.save(tmp)
            atomic_finalize(tmp, out_path)
        except BaseException:
            try:
                if os.path.exists(tmp):
                    os.remove(tmp)
            except OSError:
                pass
            raise
        return result

    # --- 子步骤 ---------------------------------------------------------- #

    def _insert_section(self, document, spec: SectionSpec) -> None:
        """插入新章节：标题段 + 正文多段。position/anchor 决定插入位置。"""
        anchor_para = self._locate_anchor(document, spec)
        title = (spec.title or "").strip()
        bodies = self._split_body(spec.body or "")

        if anchor_para is not None:
            # 在锚点段前依次插入（标题、各正文段），保持顺序。
            heading = anchor_para.insert_paragraph_before(title)
            self._make_heading(document, heading)
            for text in bodies:
                anchor_para.insert_paragraph_before(text)
        else:
            # 无锚点（position=end 或空文档）：追加到文末。
            self._make_heading(document, document.add_paragraph(title))
            for text in bodies:
                document.add_paragraph(text)

    def _locate_anchor(self, document, spec: SectionSpec):
        """定位插入锚点段落；返回段落对象或 None（None → 追加到末尾）。"""
        paragraphs = document.paragraphs
        if spec.anchor:
            for para in paragraphs:
                if spec.anchor in (para.text or ""):
                    return para
            return None
        if spec.position == "start":
            # 插到第一个非空 body 段落前（跳过纯空段）。
            for para in paragraphs:
                if (para.text or "").strip():
                    return para
            return paragraphs[0] if paragraphs else None
        return None  # position == "end" → 追加到末尾

    @staticmethod
    def _make_heading(document, paragraph) -> None:
        """把段落设为标题样式（best-effort，样式缺失则保持普通段落）。"""
        try:
            paragraph.style = document.styles["Heading 1"]
        except Exception:  # noqa: BLE001 - 无该样式则不强设，仍为新增段落
            pass

    @staticmethod
    def _split_body(body: str) -> list[str]:
        """把正文按空行切分为多段（去空段）。"""
        blocks = re.split(r"\n[ \t]*\n", body or "")
        return [b.strip() for b in blocks if b.strip()]

    def _append_references(self, document, entries: list[str], heading: str) -> int:
        """文末追加唯一一份参考文献块（已存在同名标题则不重复插标题）。"""
        from paper_agent.export.typesetting import format_reference_paragraph

        if not self._has_reference_heading(document):
            style = self._ensure_reference_style(document)
            head = document.add_paragraph(heading)
            self._make_heading(document, head)
            if style is not None:
                # 标题用 Heading，条目用参考文献受保护样式（见下）。
                pass

        style = self._ensure_reference_style(document)
        count = 0
        for i, entry in enumerate(entries, start=1):
            text = (entry or "").strip().replace("\n", " ")
            if not text:
                continue
            para = document.add_paragraph(f"{i}. {text}")
            if style is not None:
                para.style = style
            format_reference_paragraph(para)
            count += 1
        return count

    @staticmethod
    def _has_reference_heading(document) -> bool:
        """文档中是否已存在参考文献标题（去重）。"""
        for para in document.paragraphs:
            text = (para.text or "").strip().lower()
            if not text or len(text) > 20:
                continue
            if any(kw in text for kw in _REF_HEADING_KW):
                return True
        return False

    @staticmethod
    def _ensure_reference_style(document):
        """确保存在受保护的「参考文献」段落样式并返回（失败返回 None）。"""
        try:
            from docx.enum.style import WD_STYLE_TYPE
        except Exception:  # noqa: BLE001
            return None
        name = "参考文献"
        styles = document.styles
        for style in styles:
            if style.name == name:
                return style
        try:
            style = styles.add_style(name, WD_STYLE_TYPE.PARAGRAPH)
            style.base_style = styles["Normal"]
            return style
        except Exception:  # noqa: BLE001 - 样式创建失败退化为无专用样式
            return None

    @staticmethod
    def _preserved(pre: dict, post: dict, result: AugmentResult) -> bool:
        """结构计数只增不减 + 原标题集合 ⊆ 产物标题集合（Preservation_Check）。"""
        for key in ("paragraphs", "tables", "drawings", "footnote_refs"):
            if int(post.get(key, 0)) < int(pre.get(key, 0)):
                result.notes.append(f"结构计数下降：{key}")
                return False
        pre_head = set(pre.get("headings", ()) or ())
        post_head = set(post.get("headings", ()) or ())
        if not pre_head.issubset(post_head):
            result.notes.append("原有标题在产物中缺失")
            return False
        return True

    @staticmethod
    def _copy_original(in_path: str, out_path: str) -> None:
        """失败回退：把原稿复制到产物路径（保证不交付破坏性文件）。"""
        import shutil

        try:
            shutil.copyfile(in_path, out_path)
        except OSError:
            pass


def _preservation_check_docx(original_path: str, produced_path: str) -> tuple[bool, str]:
    """比对产物 docx 相对原 docx 是否结构无损（计数只增不减 + 原标题子集保留）。

    供 ``run_python`` 的 docx 微操复用（单一实现，与 :class:`InplaceDocxAugmenter` 同口径）。
    返回 ``(是否无损, 原因)``；读取失败按未通过处理（保守，不交付破坏性产物）。
    """
    try:
        import docx  # noqa: WPS433

        from paper_agent.export.docx_structural import structural_fields

        pre = structural_fields(docx.Document(original_path))
        post = structural_fields(docx.Document(produced_path))
    except Exception as exc:  # noqa: BLE001 - 读失败按未通过（保守）
        return False, f"无法读取 docx 做无损校验：{type(exc).__name__}"

    for key in ("paragraphs", "tables", "drawings", "footnote_refs"):
        if int(post.get(key, 0)) < int(pre.get(key, 0)):
            return False, f"结构计数下降：{key}"
    pre_head = set(pre.get("headings", ()) or ())
    post_head = set(post.get("headings", ()) or ())
    if not pre_head.issubset(post_head):
        return False, "原有标题在产物中缺失"
    return True, ""


__all__ = [
    "SectionSpec",
    "AugmentResult",
    "InplaceLatexAugmenter",
    "InplaceDocxAugmenter",
    "_preservation_check_docx",
]
