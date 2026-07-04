"""DOCX 结构签名与结构 diff（供 in-place 润色守卫与 Format_Gate 复用，单一真相源）。

此前 `docx_inplace` 自带一套结构签名/part-SHA 逻辑，而 `Format_Gate` 对 docx 只做
「pandoc 能否解析」的**语法级伪校验**——查不出「产物跟原文结构像不像」。本模块把
结构判定收敛到一处：

- ``structural_signature(document)``：body 结构指纹（段落/表格/图形/超链接/脚注引用
  计数 + 标题文本集合 + sectPr）——**不含 run 文本**，故仅改文字不改此签名；用于
  in-place 润色的 body 级快速比对。
- ``structural_part_shas(path)``：结构 part（styles/numbering/页眉页脚/脚注/批注等）的
  SHA——in-place 润色用「同序列化基线」比对做严格兜底。
- ``docx_structural_diff_check(pre_path, post_path)``：面向**任意两份 docx** 的语义级
  结构 diff（比对上述计数/标题/sectPr 数），**对序列化字节噪声鲁棒**——这正是
  Format_Gate 需要的「结构像不像」真校验（不比 part 原始字节，避免 python-docx
  重序列化造成的假差异）。
"""

from __future__ import annotations

import hashlib
import zipfile
from dataclasses import dataclass, field

# 结构型段落样式（子串、忽略大小写）：命中即视为「标题类」结构段，其文本纳入签名。
PROTECTED_STYLE_SUBSTR = (
    "heading", "title", "subtitle", "toc", "caption", "bibliography",
    "footnote", "endnote", "header", "footer", "table of", "标题", "题注",
    "目录", "页眉", "页脚", "参考文献",
)

# 结构 part 前缀：这些 part 内容不应被「只改 run.text」触碰。
STRUCTURAL_PART_PREFIXES = (
    "word/styles.xml",
    "word/numbering.xml",
    "word/settings.xml",
    "word/footnotes.xml",
    "word/endnotes.xml",
    "word/comments.xml",
    "word/header",
    "word/footer",
)


def qn_localname(tag: str) -> str:
    """从 lxml 限定标签 ``{ns}local`` 取本地名。"""
    return tag.rsplit("}", 1)[-1] if "}" in tag else tag


def style_is_protected(paragraph) -> bool:
    """段落样式是否属于「结构型」（标题/题注/目录/脚注/页眉页脚/参考文献等）。"""
    try:
        name = (paragraph.style.name or "") if paragraph.style else ""
    except Exception:  # noqa: BLE001 - 样式解析异常按「保护」处理
        return True
    low = name.lower()
    return any(s in low for s in PROTECTED_STYLE_SUBSTR)


def _count_local(document, tag_local: str) -> int:
    return sum(
        1 for el in document.element.iter() if qn_localname(el.tag) == tag_local
    )


def _el_to_str(el) -> str:
    try:
        from lxml import etree

        return etree.tostring(el).decode("utf-8", errors="ignore")
    except Exception:  # noqa: BLE001
        return ""


def structural_fields(document) -> dict:
    """文档结构的**带标签**字段（供签名与可读 diff 共用）。不含 run 文本。"""
    from docx.oxml.ns import qn

    body = document.element.body
    return {
        "paragraphs": len(body.findall(qn("w:p"))),
        "tables": len(body.findall(qn("w:tbl"))),
        "drawings": _count_local(document, "drawing"),
        "hyperlinks": _count_local(document, "hyperlink"),
        "footnote_refs": _count_local(document, "footnoteReference"),
        "headings": tuple(
            p.text for p in document.paragraphs if style_is_protected(p)
        ),
        "sectprs": tuple(
            _el_to_str(el)
            for el in document.element.iter()
            if qn_localname(el.tag) == "sectPr"
        ),
    }


def structural_signature(document) -> tuple:
    """结构签名元组（可用于等值比较）；等价于 ``structural_fields`` 的值序列。"""
    f = structural_fields(document)
    return (
        f["paragraphs"], f["tables"], f["drawings"], f["hyperlinks"],
        f["footnote_refs"], f["headings"], f["sectprs"],
    )


def structural_part_shas(docx_path: str) -> dict[str, str]:
    """读取 docx（zip）内结构 part 的 SHA-256。读失败返回空 dict（不抛异常）。"""
    shas: dict[str, str] = {}
    try:
        with zipfile.ZipFile(docx_path) as zf:
            for name in zf.namelist():
                if name.startswith(STRUCTURAL_PART_PREFIXES):
                    shas[name] = hashlib.sha256(zf.read(name)).hexdigest()
    except Exception:  # noqa: BLE001
        return {}
    return shas


@dataclass
class StructuralDiff:
    """两份 docx 的语义级结构 diff 结果。"""

    ok: bool
    reasons: list[str] = field(default_factory=list)


def docx_structural_diff_check(pre_path: str, post_path: str) -> StructuralDiff:
    """比对两份 docx 的语义级结构（计数/标题/sectPr），返回是否「结构一致」。

    对序列化字节噪声鲁棒（不比 part 原始字节）——这是 Format_Gate 想要的「产物跟
    原文结构像不像」真校验。任一读取失败时保守判定为不一致并给出原因。
    """
    try:
        from docx import Document
    except ImportError:  # pragma: no cover - 环境相关
        return StructuralDiff(ok=False, reasons=["python-docx 不可用，无法做结构 diff"])

    try:
        pre = structural_fields(Document(pre_path))
    except Exception as exc:  # noqa: BLE001
        return StructuralDiff(ok=False, reasons=[f"读取原文失败：{type(exc).__name__}"])
    try:
        post = structural_fields(Document(post_path))
    except Exception as exc:  # noqa: BLE001
        return StructuralDiff(ok=False, reasons=[f"读取产物失败：{type(exc).__name__}"])

    reasons: list[str] = []
    _labels = {
        "paragraphs": "段落数",
        "tables": "表格数",
        "drawings": "内嵌图形数",
        "hyperlinks": "超链接数",
        "footnote_refs": "脚注引用数",
    }
    for key, label in _labels.items():
        if pre[key] != post[key]:
            reasons.append(f"{label}不一致：原文 {pre[key]} → 产物 {post[key]}")
    if pre["headings"] != post["headings"]:
        reasons.append(
            f"标题结构不一致：原文 {len(pre['headings'])} 个 → 产物 "
            f"{len(post['headings'])} 个（或标题文本被改动）"
        )
    if len(pre["sectprs"]) != len(post["sectprs"]):
        reasons.append(
            f"分节数不一致：原文 {len(pre['sectprs'])} → 产物 {len(post['sectprs'])}"
        )
    return StructuralDiff(ok=not reasons, reasons=reasons)


__all__ = [
    "PROTECTED_STYLE_SUBSTR",
    "STRUCTURAL_PART_PREFIXES",
    "StructuralDiff",
    "qn_localname",
    "style_is_protected",
    "structural_fields",
    "structural_signature",
    "structural_part_shas",
    "docx_structural_diff_check",
]
