"""DOCX 正文排版应用：把 ``Typesetting`` 规格施加到已导出的 docx 文件。

对**正文段落**（跳过标题/题注/参考文献等结构型样式，复用
``docx_structural.style_is_protected``）设置行距 / 对齐 / 首行缩进 / 字体。未指定的
字段沿用文档既有默认（Req 6.5 语义）。

单位约定（``first_line_indent`` 字符串解析）：
- ``"<n>ch"`` → n 个字符宽（近似按 12pt 每字符换算）；
- ``"<n>pt"`` / 纯数字 → 磅；``"<n>cm"`` → 厘米。
``line_spacing`` 解释为**固定行距磅值**（对应「固定行距 22」这类需求）。
"""

from __future__ import annotations

from paper_agent.agent_platform.models import Typesetting
from paper_agent.export.docx_structural import style_is_protected

# 近似换算：1 个中文字符宽 ≈ 12 磅（首行缩进 "2ch" → 24pt）。
_CHAR_TO_PT = 12.0


def _alignment_enum(alignment: str):
    """把对齐字符串映射为 python-docx 的 WD_ALIGN_PARAGRAPH 枚举。"""
    from docx.enum.text import WD_ALIGN_PARAGRAPH

    return {
        "left": WD_ALIGN_PARAGRAPH.LEFT,
        "center": WD_ALIGN_PARAGRAPH.CENTER,
        "right": WD_ALIGN_PARAGRAPH.RIGHT,
        "justify": WD_ALIGN_PARAGRAPH.JUSTIFY,
    }.get(alignment)


def _parse_indent_pt(value: str) -> float | None:
    """把首行缩进字符串解析为磅值；无法解析返回 None。"""
    text = (value or "").strip().lower()
    if not text:
        return None
    try:
        if text.endswith("ch"):
            return float(text[:-2]) * _CHAR_TO_PT
        if text.endswith("pt"):
            return float(text[:-2])
        if text.endswith("cm"):
            return float(text[:-2]) * 28.3465  # 1cm ≈ 28.3465pt
        return float(text)  # 纯数字按磅
    except ValueError:
        return None


def apply_columns(docx_path: str, columns: int) -> int:
    """把 docx 所有 section 设为 ``columns`` 栏（节级排版原语），返回受影响的 section 数。

    这是"分栏"这一可组合排版原语的**单一实现**：操作 ``sectPr`` 的 ``w:cols num``。
    转换产物（``convert_document``）与就地排版（``polish_docx_inplace`` / set_typesetting）
    共用它，避免逻辑重复、保证行为一致。``columns<=1`` 视为单栏（清 num=1）。

    幂等、防御式：python-docx 不可用时抛 ``RuntimeError``（与 DocxExporter 一致）。
    """
    try:
        import docx  # noqa: WPS433
        from docx.oxml import OxmlElement
        from docx.oxml.ns import qn
    except ImportError as exc:  # pragma: no cover - 环境相关
        raise RuntimeError(
            "apply_columns 需要 python-docx，请安装：pip install '.[docx]'"
        ) from exc

    num = max(1, int(columns))
    document = docx.Document(docx_path)
    affected = 0
    for section in document.sections:
        sect_pr = section._sectPr
        cols = sect_pr.find(qn("w:cols"))
        if cols is None:
            cols = OxmlElement("w:cols")
            sect_pr.append(cols)
        cols.set(qn("w:num"), str(num))
        affected += 1
    document.save(docx_path)
    return affected


# 参考文献段落悬挂缩进磅值（约 2 字符宽，容纳 "[12] " / "12. " 这类编号前缀）。
REFERENCE_HANGING_INDENT_PT = 21.0


def format_reference_paragraph(
    paragraph,
    *,
    hanging_indent_pt: float = REFERENCE_HANGING_INDENT_PT,
    single_spacing: bool = True,
) -> None:
    """把一个参考文献段落设为学术标准格式：**悬挂缩进 + 单倍行距**（可组合排版原语）。

    悬挂缩进 = 左缩进 ``hanging_indent_pt`` + **首行负缩进**同值——使编号首行顶格、续行
    相对缩进，形成"悬挂"效果（GB/T 7714 与多数会议模板的参考文献表标准排版）。

    与正文排版原语（对齐/行距/分栏）正交：本原语只作用于传入的参考文献段落，且导出时
    参考文献段落使用受保护样式，故不会被 ``apply_typesetting`` 的正文规格覆盖。
    """
    from docx.enum.text import WD_LINE_SPACING
    from docx.shared import Pt

    fmt = paragraph.paragraph_format
    fmt.left_indent = Pt(hanging_indent_pt)
    fmt.first_line_indent = Pt(-hanging_indent_pt)  # 负首行缩进 = 悬挂缩进
    if single_spacing:
        fmt.line_spacing_rule = WD_LINE_SPACING.SINGLE
        fmt.line_spacing = 1.0


def apply_typesetting(docx_path: str, spec: Typesetting) -> int:
    """把排版规格施加到 docx 正文段落（含节级分栏），返回受影响的正文段落数。

    幂等、防御式：``spec`` 全未指定时不改动任何段落（返回 0）；python-docx 不可用
    时抛 ``RuntimeError``（与 DocxExporter 一致的可诊断失败）。分栏为节级原语，单独
    经 :func:`apply_columns` 施加，不计入正文段落数。
    """
    if spec.is_empty():
        return 0
    try:
        import docx  # noqa: WPS433
        from docx.shared import Pt
    except ImportError as exc:  # pragma: no cover - 环境相关
        raise RuntimeError(
            "apply_typesetting 需要 python-docx，请安装：pip install '.[docx]'"
        ) from exc

    # 分栏（节级原语）：先施加，独立于段落级排版。
    if spec.columns is not None:
        apply_columns(docx_path, spec.columns)

    document = docx.Document(docx_path)
    alignment_enum = _alignment_enum(spec.alignment) if spec.alignment else None
    indent_pt = _parse_indent_pt(spec.first_line_indent) if spec.first_line_indent else None

    affected = 0
    for para in document.paragraphs:
        if style_is_protected(para):
            continue  # 跳过标题/题注/参考文献等结构型段落
        _apply_to_paragraph(para, spec, alignment_enum, indent_pt, Pt)
        affected += 1

    document.save(docx_path)
    return affected


def _apply_to_paragraph(para, spec: Typesetting, alignment_enum, indent_pt, Pt) -> None:
    """对单个正文段落施加各已指定字段（未指定的跳过）。"""
    fmt = para.paragraph_format

    if alignment_enum is not None:
        para.alignment = alignment_enum

    if spec.line_spacing is not None:
        from docx.enum.text import WD_LINE_SPACING

        fmt.line_spacing = Pt(spec.line_spacing)
        fmt.line_spacing_rule = WD_LINE_SPACING.EXACTLY  # 固定行距

    if indent_pt is not None:
        fmt.first_line_indent = Pt(indent_pt)

    if spec.font:
        for run in para.runs:
            run.font.name = spec.font
            _set_east_asian_font(run, spec.font)


def _set_east_asian_font(run, font_name: str) -> None:
    """设置东亚字体（中文需单独设 w:eastAsia，否则中文仍用默认字体）。"""
    try:
        from docx.oxml.ns import qn

        rpr = run._element.get_or_add_rPr()
        rfonts = rpr.find(qn("w:rFonts"))
        if rfonts is None:
            rfonts = rpr.makeelement(qn("w:rFonts"), {})
            rpr.append(rfonts)
        rfonts.set(qn("w:eastAsia"), font_name)
    except Exception:  # noqa: BLE001 - 东亚字体设置失败不影响主流程
        pass


__all__ = [
    "apply_typesetting",
    "apply_columns",
    "format_reference_paragraph",
    "REFERENCE_HANGING_INDENT_PT",
]
