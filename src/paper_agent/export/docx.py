"""docx 导出器（Req 10 / venue-templates-figures-tables Req 1.6、5.x、6.3）。

面向社科与非技术协作场景。依赖 python-docx（可选依赖，见 [docx] extra），
惰性导入以保持核心零依赖。保留章节、图表说明与对已验证文献的引用（Req 10.6）。

本轮扩展（Task 13.1）：
- 应用 ``VenueProfile.docx_conventions``（如标题样式/层级）——best-effort，不可用即用默认。
- 图区：能定位导出目录内的 ``Figure_Asset`` 时 ``add_picture`` 内联嵌图并在其下加图题段落；
  无可定位资产则回落到 ``figure_id: caption`` 段落并发 ``DEGRADATION``
  （feature=figure_embed_docx, reason=missing_asset）。
- 表区：调用 ``TableRenderer.render_docx`` 追加原生 docx 表格（表头行 + 数据行）。
- python-docx 不可用：``_docx()`` 抛可诊断 ``RuntimeError``（在写盘前），不产半损坏文件。

本轮扩展（Task 14.1，format-pipeline-and-diff-revision Req 6.1/6.6）：
- 章节散文体改由 pandoc 一次性转换为**结构化 docx 主体**（标题/段落/列表/数学各映射
  为 docx 原生结构元素），修掉「整段塞进单段落」。做法：把全部章节拼接为组合
  Normalized_Markdown 文档（``# {title}\n\n{body}\n\n``），经
  ``PandocConverter.convert(..., target="docx", out_path=<.docx>)`` 产出主体，随后以
  python-docx **重新打开**该 docx，按 sibling 既有逻辑原生追加图、表与参考文献。
- pandoc 不可用（``probe`` 为假）或转换失败：回退到既有整段渲染器
  ``_export_fallback``，逐字节保持旧行为（Req 8.2 的 fallback 由 Task 16 细化标注）。
- ``[id] -> [n]`` 行内引用渲染沿用既有 docx 行为（docx 不受 Markdown 逐字 Req 7.6 约束）。
"""

from __future__ import annotations

import os
import re
import tempfile

from paper_agent.export.asset_paths import safe_relative_asset
from paper_agent.export.atomic_write import atomic_finalize
from paper_agent.export.base import ExportResult
from paper_agent.export.citation_closure import cited_references
from paper_agent.export.grounding import GroundingChecker
from paper_agent.export.inline_citations import render_inline_citations
from paper_agent.export.pandoc_pipeline import PandocConverter
from paper_agent.export.table_renderer import TableRenderer
from paper_agent.export.venue_registry import VenueRegistry
from paper_agent.observability.events import Event, EventKind, NullSink
from paper_agent.workspace.models import OutputFormat, PaperWorkspace


# pandoc 不可用 + fallback 策略时，与 DEGRADATION 事件一致的降级标注（Req 8.2）。
_PANDOC_DEGRADE_NOTE = "已降级：pandoc 不可用"

# pandoc 不可用 + fail_fast 策略时的错误说明（含安装指引，1–500 字符；Req 8.4）。
_PANDOC_FAIL_FAST_NOTE = (
    "docx 导出失败：pandoc 不可用且降级策略为 fail_fast。"
    "请安装 pandoc 后重试（安装指引见 https://pandoc.org/installing.html）。"
)


def _docx():
    try:
        import docx  # noqa: WPS433
    except ImportError as exc:  # pragma: no cover - 环境相关
        raise RuntimeError(
            "DocxExporter 需要 python-docx，请安装：pip install '.[docx]'"
        ) from exc
    return docx


class DocxExporter:
    format = OutputFormat.DOCX

    def __init__(
        self,
        sink=None,
        pandoc: PandocConverter | None = None,
        *,
        pandoc_degrade_strategy: str = "fallback",
        pandoc_probe_timeout: float = 5.0,
    ) -> None:
        # 默认 NullSink，保持 get_exporter() 的无参构造可用。
        self._sink = sink if sink is not None else NullSink()
        # 可注入 PandocConverter；缺省真实探测。probe 结果由其内部缓存，仅探测一次。
        self._pandoc = pandoc if pandoc is not None else PandocConverter()
        # pandoc 不可用时的降级策略（依赖注入，不在 exporter 内读全局 config；Req 8.2/8.4）：
        # "fallback"（默认）→ 回退整段渲染器并一致标注「已降级：pandoc 不可用」；
        # "fail_fast" → 以含 pandoc 安装指引的 1–500 字符错误说明返回空产物（不写文件）。
        # 非法值按 "fallback" 处理（合法取值集合由装配层 config 校验，见任务 1.2）。
        self._pandoc_degrade_strategy = pandoc_degrade_strategy
        # 探测 pandoc 可用性的超时秒数（Req 8.1）。
        self._pandoc_probe_timeout = pandoc_probe_timeout

    def export(self, ws: PaperWorkspace, out_dir: str) -> ExportResult:
        # 每次导出探测一次 pandoc 可用性（Req 6.1/8.1）。不可用 + fail_fast：不写任何
        # 文件，返回带 pandoc 安装指引的错误标注（Req 8.4），早于 python-docx 依赖解析。
        pandoc_available = self._pandoc.probe(timeout=self._pandoc_probe_timeout)
        if not pandoc_available and self._pandoc_degrade_strategy == "fail_fast":
            return ExportResult(
                output_format=self.format, files=[], notes=[_PANDOC_FAIL_FAST_NOTE]
            )

        # 解析可选依赖：不可用时在此抛可诊断 RuntimeError，
        # 早于 makedirs / document.save / pandoc 写盘，保证不产出半损坏 .docx（Req 5.3）。
        docx = _docx()

        # pandoc 可用 → 用 pandoc 产出结构化 docx 主体再追加原生图/表/refs（Req 6.6）；
        # 不可用（fallback 策略）→ 回退既有整段渲染器，并一致标注降级并发事件（Req 8.2）。
        if pandoc_available:
            return self._export_pandoc(ws, out_dir, docx)
        result = self._export_fallback(ws, out_dir, docx)
        result.notes.append(_PANDOC_DEGRADE_NOTE)
        self._emit_pandoc_degradation()
        return result

    def _emit_pandoc_degradation(self) -> None:
        """pandoc 不可用（fallback 策略）时发一条 DEGRADATION 事件（sink 异常不影响导出）。"""
        try:
            self._sink.emit(
                Event(
                    kind=EventKind.DEGRADATION,
                    message="pandoc 不可用，已回退整段渲染器",
                    data={
                        "feature": "pandoc",
                        "reason": "unavailable",
                        "output_format": self.format.value,
                    },
                )
            )
        except Exception:  # noqa: BLE001 - 可观测性不得中断主流程
            pass

    # ------------------------------------------------------------------ #
    # pandoc 路径：结构化正文 + 原生图/表/refs 追加
    # ------------------------------------------------------------------ #

    def _export_pandoc(self, ws: PaperWorkspace, out_dir: str, docx) -> ExportResult:
        os.makedirs(out_dir, exist_ok=True)

        notes: list[str] = []
        conventions = self._resolve_conventions(ws)

        # 引用闭合：编号与参考文献表只覆盖被正文实际引用的文献（保持既定顺序）。
        refs = cited_references(ws)
        ref_index = {r.id: i + 1 for i, r in enumerate(refs)}
        id_to_token = {rid: f"[{n}]" for rid, n in ref_index.items()}

        combined_md = self._build_combined_markdown(ws, id_to_token)

        path = os.path.join(out_dir, f"{ws.workspace_id}.docx")

        # pandoc 一次性把组合 Markdown 转为结构化 docx 主体（标题/段落/列表/数学
        # 各映射为 docx 原生结构元素）。转换失败则回退到手写整段渲染器，避免半损坏产物。
        result = self._pandoc.convert(combined_md, target="docx", out_path=path)
        if not result.ok:
            notes.append(
                "pandoc 转换 docx 主体失败，已回退整段渲染器"
                f"（exit_code={result.exit_code}）"
            )
            return self._export_fallback(ws, out_dir, docx)

        # 重新打开 pandoc 产出的 docx，按 sibling 既有逻辑原生追加图/表/参考文献。
        document = docx.Document(path)

        # 图区：优先内联嵌图（有可定位资产），否则回落文字图题段落并降级。
        if ws.figures:
            self._add_heading(document, "图表", 1, conventions)
            for fig in ws.figures:
                self._render_figure(document, ws, out_dir, fig, notes)

        # 表区：追加原生 docx 表格（放在图之后、参考文献之前）。
        TableRenderer(GroundingChecker(ws.artifact), self._sink).render_docx(
            ws.artifact, document
        )

        if refs:
            self._add_heading(document, "参考文献", 1, conventions)
            self._add_references(document, refs, ref_index)

        self._atomic_save(document, path)
        return ExportResult(output_format=self.format, files=[path], notes=notes)

    @staticmethod
    def _ensure_reference_style(document):
        """确保存在名为「参考文献」的段落样式并返回它（失败返回 None）。

        参考文献段落套用此样式后即被 ``style_is_protected`` 识别为结构段，故正文排版
        （``apply_typesetting`` 的对齐/行距/首行缩进）不会覆盖其悬挂缩进+单倍行距。
        """
        from docx.enum.style import WD_STYLE_TYPE

        name = "参考文献"
        styles = document.styles
        for style in styles:
            if style.name == name:
                return style
        try:
            style = styles.add_style(name, WD_STYLE_TYPE.PARAGRAPH)
            style.base_style = styles["Normal"]
            return style
        except Exception:  # noqa: BLE001 - 样式创建失败不阻断导出（退化为无专用样式）
            return None

    def _add_references(self, document, refs, ref_index) -> None:
        """追加参考文献段落，套用学术标准格式：悬挂缩进 + 单倍行距（复用排版原语）。"""
        from paper_agent.export.typesetting import format_reference_paragraph

        style = self._ensure_reference_style(document)
        for r in refs:
            authors = ", ".join(r.authors)
            year = r.year if r.year is not None else "n.d."
            para = document.add_paragraph(
                f"{ref_index[r.id]}. {authors} ({year}). {r.title}. "
                f"{r.source}:{r.source_id}"
            )
            if style is not None:
                para.style = style
            format_reference_paragraph(para)

    @staticmethod
    def _atomic_save(document, path: str) -> None:
        """原子保存 docx：先存到同目录临时文件再 ``os.replace``，崩溃不留半损坏文件。"""
        directory = os.path.dirname(os.path.abspath(path)) or "."
        os.makedirs(directory, exist_ok=True)
        fd, tmp = tempfile.mkstemp(dir=directory, prefix=".tmp_", suffix=".docx")
        os.close(fd)
        try:
            document.save(tmp)
            atomic_finalize(tmp, path)
        except BaseException:
            try:
                if os.path.exists(tmp):
                    os.remove(tmp)
            except OSError:
                pass
            raise

    def _build_combined_markdown(
        self, ws: PaperWorkspace, id_to_token: dict[str, str]
    ) -> str:
        """把全部章节拼接为组合 Normalized_Markdown 文档。

        各章节以 ATX 标题 + 正文拼接（``# {title}\n\n{body}\n\n``），让 pandoc 把
        标题/段落/列表/数学映射为 docx 原生结构元素（Req 6.6）。行内 ``[id]`` 沿用
        既有 docx 行为渲染为 ``[n]``；未在正文出现的已记录引用以「引用：」段落补于章节末。
        """

        blocks: list[str] = []
        for node in ws.ordered_sections():
            draft = ws.section_drafts.get(node.section_id)
            body_parts: list[str] = []
            if draft:
                content, rendered = render_inline_citations(draft.content, id_to_token)
                body_parts.append(content)
                missing = [
                    id_to_token[rid]
                    for rid in draft.cited_reference_ids
                    if rid in id_to_token and rid not in rendered
                ]
                if missing:
                    body_parts.append("引用：" + " ".join(missing))
            body = "\n\n".join(part for part in body_parts if part)
            blocks.append(f"# {node.title}\n\n{body}\n\n")
        return "".join(blocks)

    # ------------------------------------------------------------------ #
    # 回退路径：既有整段渲染器（pandoc 不可用时保持旧行为）
    # ------------------------------------------------------------------ #

    def _export_fallback(self, ws: PaperWorkspace, out_dir: str, docx) -> ExportResult:
        os.makedirs(out_dir, exist_ok=True)
        document = docx.Document()

        notes: list[str] = []

        # 会议档案 docx 约定（best-effort 格式化，缺失不失败）。
        conventions = self._resolve_conventions(ws)

        # 引用闭合：编号与参考文献表只覆盖被正文实际引用的文献（保持既定顺序）。
        refs = cited_references(ws)
        ref_index = {r.id: i + 1 for i, r in enumerate(refs)}
        id_to_token = {rid: f"[{n}]" for rid, n in ref_index.items()}

        for node in ws.ordered_sections():
            draft = ws.section_drafts.get(node.section_id)
            self._add_heading(document, node.title, 1, conventions)
            if draft:
                # 行内 [id] -> [n]（#2）；未在正文出现的已记录引用走章节末回退。
                content, rendered = render_inline_citations(draft.content, id_to_token)
                # 按段落切分渲染（而非整章塞进单段落），无 pandoc 时也保有段落结构。
                self._add_body_paragraphs(document, content)
                missing = [
                    id_to_token[rid]
                    for rid in draft.cited_reference_ids
                    if rid in id_to_token and rid not in rendered
                ]
                if missing:
                    document.add_paragraph("引用：" + " ".join(missing))

        # 图区：优先内联嵌图（有可定位资产），否则回落文字图题段落并降级。
        if ws.figures:
            self._add_heading(document, "图表", 1, conventions)
            for fig in ws.figures:
                self._render_figure(document, ws, out_dir, fig, notes)

        # 表区：追加原生 docx 表格（放在图之后、参考文献之前）。
        # TableRenderer 无数据/全空 stats 时自行发 no_data 降级并返回 0。
        TableRenderer(GroundingChecker(ws.artifact), self._sink).render_docx(
            ws.artifact, document
        )

        if refs:
            self._add_heading(document, "参考文献", 1, conventions)
            self._add_references(document, refs, ref_index)

        path = os.path.join(out_dir, f"{ws.workspace_id}.docx")
        self._atomic_save(document, path)
        return ExportResult(output_format=self.format, files=[path], notes=notes)

    # ------------------------------------------------------------------ #
    # 内部：会议约定与标题
    # ------------------------------------------------------------------ #

    @staticmethod
    def _resolve_conventions(ws: PaperWorkspace) -> dict:
        """解析 venue 的 docx_conventions；无档案/无约定即返回空 dict（用默认）。"""
        venue_id = "default"
        try:
            venue_id = ws.profile.get("venue_id", "default") or "default"
        except AttributeError:
            venue_id = "default"
        profile = VenueRegistry().resolve(venue_id)
        if profile is None:
            profile = VenueRegistry().resolve("default")
        conventions = getattr(profile, "docx_conventions", None) if profile else None
        return conventions if isinstance(conventions, dict) else {}

    @staticmethod
    def _add_body_paragraphs(document, content: str) -> None:
        """把章节正文按段落渲染为多个 docx 段落（无 pandoc 时的结构化回退）。

        - 以空行切分为段落块；
        - 每个块内：PDF 硬换行导致的行内断行做合并（中文字符间的换行直接去掉、
          其余换行转为空格），使一段连贯文字成为一个段落；
        - 形如 ``- ``/``* ``/``1. `` 的列表块逐行渲染为列表项段落。
        """
        text = content or ""
        if not text.strip():
            document.add_paragraph("")
            return

        blocks = re.split(r"\n[ \t]*\n", text)
        for block in blocks:
            block = block.strip("\n")
            if not block.strip():
                continue
            lines = [ln for ln in block.splitlines() if ln.strip()]
            if lines and all(
                re.match(r"^\s*(?:[-*·•]|\d+[.、)])\s+", ln) for ln in lines
            ):
                # 列表块：逐行成项。
                for ln in lines:
                    item = re.sub(r"^\s*(?:[-*·•]|\d+[.、)])\s+", "", ln.strip())
                    para = document.add_paragraph(item)
                    try:
                        para.style = document.styles["List Bullet"]
                    except Exception:  # noqa: BLE001 - 无该样式则用普通段落
                        pass
                continue
            # 普通段落：先去掉中文字符间的硬换行，再把其余换行并为空格。
            merged = re.sub(
                r"(?<=[\u4e00-\u9fff])[ \t]*\n[ \t]*(?=[\u4e00-\u9fff])", "", block
            )
            merged = re.sub(r"[ \t]*\n[ \t]*", " ", merged).strip()
            if merged:
                document.add_paragraph(merged)

    def _add_heading(self, document, text: str, level: int, conventions: dict):
        """按会议约定应用标题样式；不可用时回落到 python-docx 默认 add_heading。"""
        # 约定可覆盖基础标题样式名与层级（best-effort）。
        base_style = conventions.get("heading_style") if conventions else None
        eff_level = level
        try:
            level_override = conventions.get("section_heading_level") if conventions else None
            if isinstance(level_override, int) and level_override > 0:
                eff_level = level_override
        except AttributeError:
            pass
        if base_style:
            try:
                return document.add_paragraph(text, style=f"{base_style} {eff_level}")
            except Exception:  # noqa: BLE001 - 样式缺失即回落到默认标题
                pass
        return document.add_heading(text, level=eff_level)

    # ------------------------------------------------------------------ #
    # 内部：图渲染
    # ------------------------------------------------------------------ #

    def _render_figure(self, document, ws: PaperWorkspace, out_dir: str, fig, notes: list[str]) -> None:
        """嵌入单张图：有可定位资产则 add_picture + 图题；否则回落文字段落并降级。"""
        rel = safe_relative_asset(out_dir, fig.data_ref)
        if rel:
            abs_path = os.path.abspath(os.path.join(out_dir, rel))
            if os.path.isfile(abs_path):
                try:
                    document.add_picture(abs_path)
                    # 图题紧随图片之下，保证图片与图题一一对应（Property 14）。
                    document.add_paragraph(f"{fig.figure_id}: {fig.caption}")
                    return
                except Exception:  # noqa: BLE001 - 图片无法解析即回落文字图题
                    pass
        # 回落：仅输出 figure_id + 图题文字段落，并记降级（Req 5.2）。
        document.add_paragraph(f"{fig.figure_id}: {fig.caption}")
        self._sink.emit(
            Event(
                kind=EventKind.DEGRADATION,
                message="缺少可定位图资产，回落为文字图题段落",
                data={
                    "feature": "figure_embed_docx",
                    "reason": "missing_asset",
                    "figure_id": fig.figure_id,
                },
            )
        )
        notes.append(f"图 {fig.figure_id}：缺少可定位资产，已回落为文字图题（missing_asset）")
