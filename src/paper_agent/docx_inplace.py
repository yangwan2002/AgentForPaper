"""DOCX 原地润色（in-place source polish）——对标 ``latex_inplace``。

**背景（P0 头号风险修复）**：既有 `DocxExporter` 从不打开用户原 .docx，而是从
`section_drafts` 用 pandoc 从零重建——用户的字体/样式/编号/页眉页脚/TOC/交叉引用/
修订痕迹/脚注/批注/原图/原表**每次全丢**。本模块把用户的 .docx **当作真相**：
用 python-docx 打开原文，**只重写正文散文段落的 ``run.text``**，其余 OOXML 结构
（sectPr / styles / numbering / headerReference / w:tbl / w:drawing / w:hyperlink /
脚注 / 批注 / 修订）因为从不被 re-emit 而**自然逐字保留**。

保守策略（宁可少改也不破坏格式/结构）：
- **只碰正文段落**；跳过表格内段落、页眉页脚（表/页眉页脚整体保留）。
- **跳过结构型段落样式**：Heading / Title / TOC / Caption / Bibliography / 脚注 等。
- **跳过含特殊元素的段落**：超链接、域（field）、脚注/尾注/批注引用、内嵌图形/对象。
  （这些 run 本就不在 ``paragraph.runs`` 或含特殊子元素，改动风险高，一律保留。）
- **只润色格式同质段落**：段内所有非空 run 的粗体/斜体/下划线/字体/字号一致时，
  才把润色文本写回**首个 run**、清空其余 run（同质→格式观感不变）；异质段落跳过，
  避免丢失段内局部加粗等格式。
- **确定性保真守卫**：复用 ``polish_guards``——引用标注/数字多重集合恒等、长度浮动
  受限；任一破坏即丢弃该段润色、保留原文。
- **文档级结构 diff 闸 + 回滚**：写盘前比对 pre/post 的结构签名（段落/表格/图形/
  超链接/脚注计数、标题文本集合、sectPr）——任一不等则**整档回滚**（输出原文副本），
  绝不产出结构被破坏的文件。
- **永不覆盖用户原文件**：产物写到独立 out_path；输入文件全程只读。

Mock provider（``is_mock=True``）下整体 no-op：产物在结构与文字上等同输入。
"""

from __future__ import annotations

import shutil
from dataclasses import dataclass, field

from paper_agent.export.atomic_write import atomic_finalize
from paper_agent.export.docx_structural import (
    qn_localname as _qn_localname,
    structural_part_shas as _structural_part_shas,
    structural_signature as _structural_signature,
    style_is_protected as _style_is_protected,
)
from paper_agent.inplace_core import ProsePolishGuard, polish_fragment
from paper_agent.prompts import templates
from paper_agent.providers.llm.base import LLMProvider
from paper_agent.tools import polish_guards

# DOCX 散文守卫：内容（引用/数字）+ 长度须保持（无 LaTeX 结构约束）。
_DOCX_GUARD = ProsePolishGuard(
    [polish_guards.content_preserved, polish_guards.length_ratio_ok]
)

# 段落内一旦出现这些 OOXML 元素即跳过（本地名，配 w 命名空间比对）。
# 含 ``ins``/``del``：修订痕迹（tracked changes）——``paragraph.runs`` 不列出 w:ins/
# w:del 内嵌 run，若润色后写回 runs[0] 会与保留的 ins/del run 并存、造成「接受修订」
# 后文本 = 润色版 + 旧删除版拼接的语义错乱，故含修订痕迹的段落一律跳过。
# 结构签名 / part-SHA / 标题样式判定统一复用 ``export.docx_structural``（单一真相源）。
_SPECIAL_LOCAL_NAMES = frozenset(
    {
        "hyperlink", "fldSimple", "instrText",
        "footnoteReference", "endnoteReference", "commentReference",
        "drawing", "object", "pict",
        "ins", "del",
    }
)

# 一个段落至少含这么多字母/汉字才值得润色（否则纯符号/空白，跳过）。
_MIN_PARA_LETTERS = 15


@dataclass
class InplaceDocxResult:
    """DOCX 原地润色结果。"""

    out_path: str
    total_prose_paragraphs: int = 0     # 判定为可润色的正文段落数
    polished_paragraphs: int = 0        # 实际润色替换的段落数
    rejected_by_guard: int = 0          # 因保真守卫拦截保留原文的段落数
    rolled_back: bool = False           # 结构 diff 闸失败 → 整档回滚为原文副本
    notes: list[str] = field(default_factory=list)


def _letters(text: str) -> int:
    return sum(
        1 for ch in text
        if ch.isalpha() or "\u4e00" <= ch <= "\u9fff"
    )


def _paragraph_has_special(paragraph) -> bool:
    """段落内是否含超链接/域/脚注/批注/图形等特殊元素（含则跳过）。"""
    for el in paragraph._p.iter():
        if _qn_localname(el.tag) in _SPECIAL_LOCAL_NAMES:
            return True
    return False


def _runs_homogeneous(runs) -> bool:
    """段内所有非空 run 的关键格式属性是否一致（同质才安全把文本并入首 run）。"""
    sig = None
    for r in runs:
        if not r.text:
            continue
        font = r.font
        size = font.size
        cur = (
            bool(r.bold), bool(r.italic), bool(r.underline),
            font.name or "", size.pt if size is not None else 0.0,
        )
        if sig is None:
            sig = cur
        elif cur != sig:
            return False
    return True


def _should_polish(paragraph) -> bool:
    runs = paragraph.runs
    if not runs:
        return False
    text = "".join(r.text for r in runs)
    if _letters(text) < _MIN_PARA_LETTERS:
        return False
    if _style_is_protected(paragraph):
        return False
    if _paragraph_has_special(paragraph):
        return False
    if not _runs_homogeneous(runs):
        return False
    return True


def _safe_remove(path: str) -> None:
    import os

    try:
        if path and os.path.exists(path):
            os.remove(path)
    except OSError:
        pass


def _count_table_prose_paragraphs(document) -> int:
    """统计表格单元格内「够长的散文」段落数（这些段落被原地润色跳过，需向用户报告）。"""
    n = 0
    for table in document.tables:
        for row in table.rows:
            for cell in row.cells:
                for p in cell.paragraphs:
                    text = "".join(r.text for r in p.runs)
                    if _letters(text) >= _MIN_PARA_LETTERS:
                        n += 1
    return n


class InplaceDocxPolisher:
    """DOCX 原地润色器：只改正文散文 run.text，保 OOXML 结构，结构 diff 闸兜底。"""

    def __init__(self, llm: LLMProvider, *, is_mock: bool = False) -> None:
        self._llm = llm
        self._is_mock = is_mock

    def polish(self, in_path: str, out_path: str) -> InplaceDocxResult:
        """润色 ``in_path`` 的 .docx，产物写到 ``out_path``（输入只读，永不覆盖）。

        Mock provider 下 no-op：直接复制原文到 out_path（结构与文字均等同输入）。
        """
        try:
            import docx  # noqa: WPS433 - 可选依赖，惰性导入
        except ImportError as exc:  # pragma: no cover - 环境相关
            raise RuntimeError(
                "DOCX 原地润色需要 python-docx，请安装：pip install '.[docx]'"
            ) from exc

        if self._is_mock:
            shutil.copyfile(in_path, out_path)
            return InplaceDocxResult(
                out_path=out_path,
                notes=["Mock provider：DOCX 原地润色 no-op（复制原文）。"],
            )

        document = docx.Document(in_path)
        # pre 结构签名（body 计数 + 标题文本 + sectPr）。
        pre_sig = _structural_signature(document)
        # 结构 part 基线 SHA：把**未修改**的文档经同一序列化路径存到临时文件再哈希，
        # 以消除 python-docx 重序列化的字节噪声（真实差异才会被后续比对捕获）。
        base_tmp = self._save_tmp(document)
        pre_parts = _structural_part_shas(base_tmp)
        _safe_remove(base_tmp)

        table_prose = _count_table_prose_paragraphs(document)
        total = 0
        polished = 0
        rejected = 0
        for paragraph in document.paragraphs:
            if not _should_polish(paragraph):
                continue
            total += 1
            runs = paragraph.runs
            original = "".join(r.text for r in runs)
            new_text = self._polish_paragraph(original)
            if new_text is None:
                rejected += 1
                continue
            if new_text == original:
                continue
            # 同质段落：文本并入首个 run，清空其余 run（格式观感不变）。
            runs[0].text = new_text
            for r in runs[1:]:
                r.text = ""
            polished += 1

        # 修改后先存到临时文件（建在 out_path 同目录，保证 os.replace 不跨盘），
        # 再据其真实 zip 比对结构签名与结构 part SHA。
        final_tmp = self._save_tmp(document, near_path=out_path)
        post_sig = _structural_signature(document)
        post_parts = _structural_part_shas(final_tmp)

        table_note = (
            [f"注意：跳过 {table_prose} 个表格内段落（表格整体保留，未润色其文字）。"]
            if table_prose
            else []
        )

        # 文档级结构 diff 闸：body 签名 + 结构 part SHA 均须与基线一致，否则整档回滚。
        if post_sig != pre_sig or post_parts != pre_parts:
            _safe_remove(final_tmp)
            shutil.copyfile(in_path, out_path)
            return InplaceDocxResult(
                out_path=out_path,
                total_prose_paragraphs=total,
                polished_paragraphs=0,
                rejected_by_guard=rejected,
                rolled_back=True,
                notes=["结构 diff 闸失败：润色改变了文档结构，已整档回滚为原文。"]
                + table_note,
            )

        # 结构一致：把已写好的临时文件原子替换为产物（避免再存一次）。
        atomic_finalize(final_tmp, out_path)
        return InplaceDocxResult(
            out_path=out_path,
            total_prose_paragraphs=total,
            polished_paragraphs=polished,
            rejected_by_guard=rejected,
            notes=[
                f"DOCX 原地润色：正文散文段 {total} 个，润色 {polished} 个，"
                f"守卫拦截保留原文 {rejected} 个；表格/图/页眉页脚/样式/脚注/修订等结构逐字保留。"
            ]
            + table_note,
        )

    def _polish_paragraph(self, text: str) -> str | None:
        """润色单个段落文本：委托共享核心（整段送 LLM、守卫兜底）。"""
        return polish_fragment(
            self._llm,
            lambda core: templates.polish_plain_prose(fragment=core),
            text,
            _DOCX_GUARD,
            preserve_edges=False,
        )

    @staticmethod
    def _save_tmp(document, near_path: str | None = None) -> str:
        """把 document 存到临时 .docx 并返回路径。

        ``near_path`` 提供时把临时文件建在其所在目录，使后续 ``os.replace`` 原子替换
        不跨盘符（Windows 跨盘 replace 会失败）；否则用系统临时目录（仅供哈希后删除）。
        """
        import os
        import tempfile

        if near_path:
            directory = os.path.dirname(os.path.abspath(near_path)) or "."
            os.makedirs(directory, exist_ok=True)
        else:
            directory = None
        fd, tmp = tempfile.mkstemp(prefix=".paperagent_", suffix=".docx", dir=directory)
        os.close(fd)
        try:
            document.save(tmp)
        except BaseException:
            _safe_remove(tmp)
            raise
        return tmp


__all__ = ["InplaceDocxPolisher", "InplaceDocxResult"]
