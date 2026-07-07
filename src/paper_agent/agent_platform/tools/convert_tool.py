"""convert_document 工具：用户原文件的**跨格式直转**（保公式/结构，可选双栏）。

问题背景：把用户的 `.tex` 转 docx 时，走 import_draft（抽纯文本）→ 当 Markdown → 重建
docx 的路径会**丢公式、乱结构**（LaTeX 数学被当普通文字）。正确做法是让 pandoc 以
**LaTeX 为输入格式直转** docx（``pandoc -f latex -t docx``）——公式变 Word 原生公式、
章节结构保留。

本工具就是这条"源文件直转"路径：拿用户原始 `.tex`/`.docx`/`.md` 当真相，用
:meth:`PandocConverter.convert_file` 直转目标格式；docx 目标可选叠加**双栏**（Word
sectPr 分栏）与已保存的排版规格。是"用户给 A 格式、要 B 格式"这类纯转换任务的正确工具，
不走重建。
"""

from __future__ import annotations

import os
from dataclasses import dataclass, field

from paper_agent.agent_platform.models import Typesetting
from paper_agent.agent_platform.tools.context import ToolContext
from paper_agent.export.latex_normalize import normalize_latex_for_pandoc
from paper_agent.export.pandoc_pipeline import PandocConverter
from paper_agent.tools.registry import ToolRegistry


@dataclass
class ConvertOutcome:
    """跨格式转换的结构化结果（供工具与工作流复用）。

    ``notes`` 是可直接拼接成用户可读文本的片段列表（首片无前导空格、后续含前导
    空格，``"".join(notes)`` 即完整说明）；失败时 ``ok=False`` 且 ``error`` 非空。
    """

    ok: bool
    files: list[str] = field(default_factory=list)
    notes: list[str] = field(default_factory=list)
    error: str = ""

    def message(self) -> str:
        """人可读结论：失败给 error，成功给拼接后的 notes。"""
        return self.error if not self.ok else "".join(self.notes)

# 源文件扩展名 → pandoc 输入格式。
_EXT_TO_FROM = {
    ".tex": "latex", ".latex": "latex",
    ".docx": "docx",
    ".md": "markdown", ".markdown": "markdown",
}
# 目标格式 → (pandoc 输出格式, 产物扩展名)。
_TO_FORMAT = {
    "docx": ("docx", ".docx"),
    "latex": ("latex", ".tex"),
    "markdown": ("markdown", ".md"),
}

_SCHEMA = {
    "type": "object",
    "properties": {
        "to_format": {
            "type": "string",
            "enum": sorted(_TO_FORMAT),
            "description": "目标格式：docx / latex / markdown。",
        },
        "path": {
            "type": "string",
            "description": "原文件绝对路径；省略则用本会话已导入的原文件。",
        },
        "two_column": {
            "type": "boolean",
            "description": "仅对 docx 目标有效：是否设为双栏排版（小论文常用）。默认 false。",
        },
        "three_line_table": {
            "type": "boolean",
            "description": (
                "仅对 docx 目标有效：是否把表格套用三线表（booktabs）样式——学术论文"
                "标准表格。默认 true。"
            ),
        },
    },
    "required": ["to_format"],
}

_DESCRIPTION = (
    "把用户的原文件**跨格式直转**为目标格式（如 .tex → docx），用 pandoc 直接转换，"
    "**保留公式与章节结构**（LaTeX 公式会转成 Word 原生公式）。docx 目标可选设双栏、"
    "并套用 set_typesetting 的排版。当用户要求「把 X 格式转成 Y 格式」这类纯转换时用本"
    "工具；不要用 import_draft + export_paper 重建（那会丢公式、乱结构）。"
)


def _resolve_source(ctx: ToolContext, path: str | None) -> tuple[str | None, str | None]:
    """解析源文件路径与 pandoc 输入格式；无法解析返回 (None, None)。"""
    candidate = (path or "").strip().strip('"').strip("'")
    if not candidate:
        profile = getattr(ctx.workspace, "profile", None) or {}
        candidate = profile.get("source_document_path", "")
    if not candidate or not os.path.isfile(candidate):
        return None, None
    ext = os.path.splitext(candidate)[1].lower()
    from_format = _EXT_TO_FROM.get(ext)
    return (candidate, from_format) if from_format else (candidate, None)


def _set_two_columns(docx_path: str) -> None:
    """把 docx 所有 section 设为双栏。委托给分栏排版原语（单一实现，避免重复）。"""
    from paper_agent.export.typesetting import apply_columns

    apply_columns(docx_path, 2)


def _fix_table_widths(docx_path: str) -> None:
    """把 docx 表格的固定列宽清成自适应（autofit），消除窄栏里逐字符压缩换行。

    pandoc 从 LaTeX 转出的表格常带固定 ``tcW`` / ``tblW`` 宽度；放进双栏窄栏时 Word
    硬按固定宽度塞，导致每列被压到最小、内容逐字符换行。改为 autofit（表格与单元格
    宽度设 ``auto`` + ``tblLayout=autofit``）后，Word 按内容重算列宽。
    """
    import docx  # noqa: WPS433
    from docx.enum.table import WD_TABLE_ALIGNMENT
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn

    document = docx.Document(docx_path)
    for table in document.tables:
        table.allow_autofit = True
        # 表格整体居中（学术论文惯例；pandoc 默认靠左，此处纠正）。
        try:
            table.alignment = WD_TABLE_ALIGNMENT.CENTER
        except Exception:  # noqa: BLE001 - 对齐设置失败不影响宽度修复
            pass
        tbl = table._tbl
        tbl_pr = tbl.tblPr
        # tblLayout = autofit（否则 Word 用固定布局，严格按 gridCol 比例压缩）。
        layout = tbl_pr.find(qn("w:tblLayout"))
        if layout is None:
            layout = OxmlElement("w:tblLayout")
            tbl_pr.append(layout)
        layout.set(qn("w:type"), "autofit")
        # 表格总宽 → auto。
        for tbl_w in tbl_pr.findall(qn("w:tblW")):
            tbl_w.set(qn("w:w"), "0")
            tbl_w.set(qn("w:type"), "auto")
        # 各单元格固定宽度 → auto（关键：消除逐字符压缩）。
        for tc_w in tbl.iter(qn("w:tcW")):
            tc_w.set(qn("w:w"), "0")
            tc_w.set(qn("w:type"), "auto")
    document.save(docx_path)


def _apply_three_line_table_style(docx_path: str) -> None:
    """给 docx 所有表格套用**三线表（booktabs）**样式：仅顶线、表头下线、底线；
    去掉所有竖线与其它横线——学术论文标准表格样式。

    边框粗细（sz 单位为 1/8 pt）：顶/底线 1.5pt（sz=12），表头下线 0.75pt（sz=6）。
    """
    import docx  # noqa: WPS433
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn

    def _border(name: str, *, val: str, sz: int = 0) -> "OxmlElement":
        el = OxmlElement(f"w:{name}")
        el.set(qn("w:val"), val)
        if val != "none":
            el.set(qn("w:sz"), str(sz))
            el.set(qn("w:space"), "0")
            el.set(qn("w:color"), "000000")
        return el

    document = docx.Document(docx_path)
    for table in document.tables:
        tbl_pr = table._tbl.tblPr
        # 重置表格边框：顶/底粗线，左右/内部竖线/内部横线全去掉。
        old = tbl_pr.find(qn("w:tblBorders"))
        if old is not None:
            tbl_pr.remove(old)
        borders = OxmlElement("w:tblBorders")
        borders.append(_border("top", val="single", sz=12))
        borders.append(_border("bottom", val="single", sz=12))
        borders.append(_border("left", val="none"))
        borders.append(_border("right", val="none"))
        borders.append(_border("insideH", val="none"))
        borders.append(_border("insideV", val="none"))
        tbl_pr.append(borders)
        # 表头行（第一行）底部加细线（\midrule）。
        if table.rows:
            for cell in table.rows[0].cells:
                tc_pr = cell._tc.get_or_add_tcPr()
                tc_borders = tc_pr.find(qn("w:tcBorders"))
                if tc_borders is None:
                    tc_borders = OxmlElement("w:tcBorders")
                    tc_pr.append(tc_borders)
                existing_bottom = tc_borders.find(qn("w:bottom"))
                if existing_bottom is not None:
                    tc_borders.remove(existing_bottom)
                tc_borders.append(_border("bottom", val="single", sz=6))
    document.save(docx_path)


# 宽表适配参数（把 LaTeX 里 \footnotesize + \tabcolsep=3pt + \extracolsep{\fill} 那套
# "挤得下"的手段补回 docx——pandoc 转换会把它们全丢掉，导致 Word 逐字符折行）。
_TABLE_FONT_PT = 9.0            # 表格字体缩到小五号（≈ LaTeX \footnotesize，正文小 1 号）
_TABLE_CELL_MARGIN_TWIPS = 40   # 单元格左右内边距（≈2pt；Word 默认 108twips 太宽，8 列吃不消）
_WIDE_TABLE_MIN_COLS = 6        # 列数 ≥ 此阈值判为"宽表"，双栏文档里让其跨栏铺满页宽


def _compact_tables(docx_path: str) -> int:
    """把所有表格紧凑化：小号字体 + 压缩单元格边距 + 占满可用宽度。

    对应补回 pandoc 丢失的 ``\\footnotesize`` / ``\\tabcolsep`` / ``\\extracolsep{\\fill}``：
    Word 默认用正文字号 + 宽边距 + 无列宽指引渲染多列表，导致每格逐字符折行。返回表格数。
    """
    import docx  # noqa: WPS433
    from docx.shared import Pt

    document = docx.Document(docx_path)
    text_width_twips = _section_text_width_twips(document)
    count = 0
    for table in document.tables:
        for row in table.rows:
            for cell in row.cells:
                for para in cell.paragraphs:
                    for run in para.runs:
                        run.font.size = Pt(_TABLE_FONT_PT)
        _set_table_cell_margins(table, _TABLE_CELL_MARGIN_TWIPS)
        # 关键修复：按内容比例给各列分配明确宽度（固定布局），而不是丢给 Word autofit
        # 乱猜——后者会把「区间/标签」列压太窄导致逐字符/逐词折行。同时对短单元格禁折行。
        _apply_content_proportional_widths(table, text_width_twips)
        count += 1
    document.save(docx_path)
    return count


# 一个显示宽度单位 ≈ 一个西文字符；CJK 记 2 单位（更接近实际占宽）。
def _display_width(text: str) -> int:
    width = 0
    for ch in text or "":
        width += 2 if ord(ch) >= 0x2E80 else 1
    return width


def _section_text_width_twips(document) -> int:
    """版心宽（页宽 − 左右页边距）换算成 twips；取不到时回退 A4 版心近似值。"""
    try:
        sec = document.sections[0]
        emu = int(sec.page_width) - int(sec.left_margin) - int(sec.right_margin)
        twips = emu // 635  # 1 twip = 635 EMU
        if 1440 <= twips <= 20000:
            return twips
    except Exception:  # noqa: BLE001 - 取不到 section 尺寸 → 回退
        pass
    return 9200  # ≈ A4 版心（21cm − 2×2.54cm）的 twips 近似


def _apply_content_proportional_widths(table, total_twips: int) -> None:
    """按各列内容长度比例分配列宽（固定布局），并对无空格短单元格禁折行。

    列权重 = 该列所有单元格显示宽度的最大值（表头+正文），据此把 ``total_twips`` 按比例
    分给各列（每列给个下限，避免某列被压没）。这样「区间 / 标签」等长内容列拿到足够宽度、
    不再逐词折行；``N`` 等短列不浪费宽度。取代 pandoc 转出后 Word autofit 的糟糕猜测。
    """
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn

    tbl = table._tbl
    grid = tbl.find(qn("w:tblGrid"))
    if grid is None:
        return
    grid_cols = grid.findall(qn("w:gridCol"))
    ncols = len(grid_cols)
    if ncols == 0:
        return

    # 各列权重 = 列内单元格显示宽度的最大值（含少量内边距余量）。
    weights = [1] * ncols
    for row in table.rows:
        cells = row.cells
        for j in range(min(ncols, len(cells))):
            w = _display_width(cells[j].text.strip()) + 2  # +2 余量
            if w > weights[j]:
                weights[j] = w
    total_weight = sum(weights) or ncols

    min_col = 500  # 每列宽度下限（twips，≈0.35cm），防某列被压没
    usable = max(total_twips, min_col * ncols)
    col_twips = [max(min_col, round(usable * wt / total_weight)) for wt in weights]

    # 固定布局 + 表宽 = 各列之和（dxa）。
    tbl_pr = tbl.tblPr
    layout = tbl_pr.find(qn("w:tblLayout"))
    if layout is None:
        layout = OxmlElement("w:tblLayout")
        tbl_pr.append(layout)
    layout.set(qn("w:type"), "fixed")
    tbl_w = tbl_pr.find(qn("w:tblW"))
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:type"), "dxa")
    tbl_w.set(qn("w:w"), str(sum(col_twips)))

    # 写 gridCol 宽度（gridCol 只有 w:w，无 type 属性）。
    for j, gc in enumerate(grid_cols):
        gc.set(qn("w:w"), str(col_twips[j]))

    # 写每个单元格的 tcW = 对应列宽；无空格短内容单元格加 noWrap（防「关联类型/UAV–Relay」折行）。
    for row in table.rows:
        cells = row.cells
        for j in range(min(ncols, len(cells))):
            tc_pr = cells[j]._tc.get_or_add_tcPr()
            tc_w = tc_pr.find(qn("w:tcW"))
            if tc_w is None:
                tc_w = OxmlElement("w:tcW")
                tc_pr.append(tc_w)
            tc_w.set(qn("w:w"), str(col_twips[j]))
            tc_w.set(qn("w:type"), "dxa")
            txt = cells[j].text.strip()
            if txt and " " not in txt and _display_width(txt) <= 14:
                if tc_pr.find(qn("w:noWrap")) is None:
                    tc_pr.append(OxmlElement("w:noWrap"))


def _set_table_cell_margins(table, twips: int) -> None:
    """设置整表默认单元格内边距（左右收紧、上下归零），模拟 LaTeX 小 ``\\tabcolsep``。"""
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn

    tbl_pr = table._tbl.tblPr
    mar = tbl_pr.find(qn("w:tblCellMar"))
    if mar is None:
        mar = OxmlElement("w:tblCellMar")
        tbl_pr.append(mar)
    for side, val in (("top", "0"), ("bottom", "0"),
                      ("left", str(twips)), ("right", str(twips))):
        el = mar.find(qn(f"w:{side}"))
        if el is None:
            el = OxmlElement(f"w:{side}")
            mar.append(el)
        el.set(qn("w:w"), val)
        el.set(qn("w:type"), "dxa")


def _span_wide_tables(docx_path: str, *, min_cols: int = _WIDE_TABLE_MIN_COLS) -> int:
    """双栏文档里，把列数 ≥ ``min_cols`` 的宽表用连续分节符裹成**单栏岛**，独占整页宽。

    等价于 LaTeX 里 ``table*`` 的"跨栏浮动"：在宽表前后各插一个带 ``sectPr`` 的空段——
    表前 sectPr(cols=2) 结束前面的双栏区，表后 sectPr(cols=1) 使"含该表的这一节"变单栏，
    全部 ``continuous`` 不产生分页。必须在所有分栏设置之后运行（否则被 apply_columns 覆盖）。
    返回被跨栏处理的宽表数。
    """
    import copy

    import docx  # noqa: WPS433
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn

    def _sectpr_para(num_cols: int, template):
        p = OxmlElement("w:p")
        p_pr = OxmlElement("w:pPr")
        p.append(p_pr)
        sect_pr = OxmlElement("w:sectPr")
        p_pr.append(sect_pr)
        sec_type = OxmlElement("w:type")
        sec_type.set(qn("w:val"), "continuous")
        sect_pr.append(sec_type)
        if template is not None:  # 沿用页面尺寸/页边距，避免新节回落到默认纸张
            for tag in ("w:pgSz", "w:pgMar"):
                src = template.find(qn(tag))
                if src is not None:
                    sect_pr.append(copy.deepcopy(src))
        cols = OxmlElement("w:cols")
        cols.set(qn("w:num"), str(num_cols))
        sect_pr.append(cols)
        return p

    document = docx.Document(docx_path)
    body = document.element.body
    body_sectpr = body.find(qn("w:sectPr"))
    spanned = 0
    for table in document.tables:
        tbl = table._tbl
        grid = tbl.find(qn("w:tblGrid"))
        ncols = len(grid.findall(qn("w:gridCol"))) if grid is not None else 0
        if ncols < min_cols:
            continue
        tbl.addprevious(_sectpr_para(2, body_sectpr))  # 关掉表前的双栏区
        tbl.addnext(_sectpr_para(1, body_sectpr))       # 含表的这一节 = 单栏铺满
        spanned += 1
    document.save(docx_path)
    return spanned


def _apply_typesetting_if_any(ctx: ToolContext, docx_path: str) -> str:
    """docx 目标：套用已保存的排版规格（set_typesetting 落定的），返回说明后缀。"""
    profile = getattr(ctx.workspace, "profile", None) or {}
    spec_data = profile.get("typesetting")
    if not spec_data:
        return ""
    spec = Typesetting.from_dict(spec_data)
    if spec.is_empty():
        return ""
    try:
        from paper_agent.export.typesetting import apply_typesetting

        apply_typesetting(docx_path, spec)
        return " 已套用排版规格。"
    except Exception as exc:  # noqa: BLE001 - 排版失败不影响转换产物
        return f"（排版应用失败：{exc}）"


def _is_locked_error(stderr: str) -> bool:
    """pandoc stderr 是否为"输出文件被占用/权限拒绝"（多为目标 docx 在 Word 中打开）。"""
    return "permission denied" in (stderr or "").lower()


def _unique_out_path(output_dir: str, stem: str, ext: str) -> str:
    """在 output_dir 里生成一个不与现有文件冲突的新产物路径（加时间戳后缀）。"""
    import time

    token = time.strftime("%H%M%S")
    candidate = os.path.join(output_dir, f"{stem}_converted_{token}{ext}")
    i = 1
    while os.path.exists(candidate):
        candidate = os.path.join(output_dir, f"{stem}_converted_{token}_{i}{ext}")
        i += 1
    return candidate


def convert_document_core(
    ctx: ToolContext, to_format: str, path: str | None = None,
    two_column: bool = False, three_line_table: bool = True,
) -> ConvertOutcome:
    """跨格式直转的**确定性核心**：解析源 → pandoc 直转 → docx 后处理（表格/双栏/排版）。

    返回结构化 :class:`ConvertOutcome`，供 ``convert_document`` 工具与 ``ConvertWorkflow``
    共享同一实现（工具只负责格式化文本、工作流只负责按固定序编排 + 诚实上报）。
    产物写新文件、原稿只读；任何一步失败经 ``ok/error`` 诚实上报，不降级重建。
    """
    to_format = (to_format or "").lower()
    if to_format not in _TO_FORMAT:
        return ConvertOutcome(
            ok=False, error=f"不支持的目标格式：{to_format}；支持 {sorted(_TO_FORMAT)}。"
        )

    src, from_format = _resolve_source(ctx, path)
    if src is None:
        return ConvertOutcome(
            ok=False,
            error="未找到可转换的源文件：请提供原文件的绝对路径，或先用 import_draft 导入。",
        )
    if from_format is None:
        return ConvertOutcome(
            ok=False,
            error=f"无法识别源文件类型（{os.path.splitext(src)[1]}），支持 .tex/.docx/.md。",
        )

    pandoc = PandocConverter()
    if not pandoc.probe():
        return ConvertOutcome(
            ok=False,
            error=(
                "转换失败：未检测到 pandoc（跨格式直转依赖它）。若已安装但不在 PATH，"
                "请设置环境变量 PANDOC_PATH 指向 pandoc.exe（如 "
                "PANDOC_PATH=D:\\path\\to\\pandoc.exe）；未安装见 "
                "https://pandoc.org/installing.html。"
            ),
        )

    out_format, out_ext = _TO_FORMAT[to_format]
    if from_format == out_format:
        return ConvertOutcome(
            ok=True, notes=[f"源文件已是 {to_format} 格式，无需转换。"]
        )

    stem = os.path.splitext(os.path.basename(src))[0]
    os.makedirs(ctx.output_dir, exist_ok=True)
    out_path = os.path.join(ctx.output_dir, f"{stem}_converted{out_ext}")

    # LaTeX 源转换前的确定性预规整（如折平 \shortstack/\makecell 单元格内换行，防表格
    # 被 pandoc 误当纯文本吐出）。在**副本**上做，原文件只读；资源路径仍指向原目录。
    notes: list[str] = []
    in_path, cleanup_path, resource_dir, prep_notes = _prepare_latex_input(
        src, from_format, ctx.output_dir, stem
    )
    notes.extend(prep_notes)

    result = pandoc.convert_file(
        in_path, out_path, from_format=from_format, to_format=out_format,
        resource_dir=resource_dir,
    )
    # 输出文件被占用（多为上次转出的 docx 仍在 Word/WPS 中打开、被 Windows 锁定）→
    # 换个不冲突的新文件名重试，并如实告知，不让用户对着 "permission denied" 猜。
    if not result.ok and _is_locked_error(result.stderr):
        alt_path = _unique_out_path(ctx.output_dir, stem, out_ext)
        result = pandoc.convert_file(
            in_path, alt_path, from_format=from_format, to_format=out_format,
            resource_dir=resource_dir,
        )
        if result.ok:
            out_path = alt_path
            notes.append(
                f"（原输出文件可能正被 Word/WPS 打开而锁定，已改存为 "
                f"{os.path.basename(alt_path)}。）"
            )
    _cleanup_temp(cleanup_path)  # 预规整临时副本用完即删（原文件从不受影响）。
    if not result.ok:
        hint = ""
        if _is_locked_error(result.stderr):
            hint = (
                "：输出文件被占用——请关闭正在打开该 docx 的 Word/WPS 后重试，"
                "或改用其它输出目录。"
            )
        return ConvertOutcome(
            ok=False,
            error=(
                f"pandoc 转换失败（exit={result.exit_code}）{hint}"
                f"{('：' + result.stderr[:300]) if not hint else ''}"
            ),
        )

    notes.insert(0, f"已直转为 {to_format}：{out_path}（公式与结构由 pandoc 保留）。")
    if to_format == "docx":
        # 先修表格列宽（清固定宽度为自适应），避免双栏窄栏里逐字符压缩。
        try:
            _fix_table_widths(out_path)
            notes.append(" 表格已设自适应列宽。")
        except Exception as exc:  # noqa: BLE001 - 表格修复失败不影响主产物
            notes.append(f"（表格宽度修复失败：{exc}）")
        if three_line_table:
            try:
                _apply_three_line_table_style(out_path)
                notes.append(" 表格已套用三线表样式。")
            except Exception as exc:  # noqa: BLE001 - 样式失败不影响主产物
                notes.append(f"（三线表样式应用失败：{exc}）")
        # 宽表适配：小号字体 + 紧凑边距 + 占满宽度（补回 pandoc 丢失的 LaTeX 表格排版），
        # 消除多列表在窄栏里逐字符折行的"难看"问题。
        try:
            _compact_tables(out_path)
            notes.append(" 表格已紧凑化（小号字体+紧凑边距+占满宽度）。")
        except Exception as exc:  # noqa: BLE001 - 紧凑化失败不影响主产物
            notes.append(f"（表格紧凑化失败：{exc}）")
        if two_column:
            try:
                _set_two_columns(out_path)
                notes.append(" 已设为双栏。")
            except Exception as exc:  # noqa: BLE001 - 双栏失败不影响主产物
                notes.append(f"（双栏设置失败：{exc}）")
        ts_note = _apply_typesetting_if_any(ctx, out_path)
        if ts_note:
            notes.append(ts_note)
        # 宽表跨栏必须在所有分栏设置（含 typesetting 里的 apply_columns）**之后**，
        # 否则新插的单栏岛 sectPr 会被 apply_columns 统一改回多栏而失效。
        if two_column:
            try:
                spanned = _span_wide_tables(out_path)
                if spanned:
                    notes.append(f" {spanned} 张宽表已跨双栏铺满页宽。")
            except Exception as exc:  # noqa: BLE001 - 跨栏失败不影响主产物
                notes.append(f"（宽表跨栏失败：{exc}）")

    ctx.session.record("convert_document", source=src, to=to_format, files=[out_path])
    # 记住源文件与本次产物到 profile：后续轮次「改双栏/调格式」等能凭记忆定位原文件，
    # 不必用户每轮重贴路径（修「转完就忘了源文件」的记忆断层）。
    try:
        ctx.repo.update(ctx.workspace, _remember_source_mutation(src, out_path))
    except Exception:  # noqa: BLE001 - 记忆写入失败不影响转换产物
        pass
    return ConvertOutcome(ok=True, files=[out_path], notes=notes)


def _prepare_latex_input(
    src: str, from_format: str | None, output_dir: str, stem: str
) -> tuple[str, str | None, str | None, list[str]]:
    """LaTeX 源转换前预规整：在副本上折平 pandoc 不友好构造，原文件只读。

    返回 ``(输入路径, 待清理临时路径 or None, 资源目录 or None, 说明片段)``：
    - 非 LaTeX、读文件失败、或无需改动 → 原样用 ``src``（临时路径为 None）。
    - 有改动 → 写规整后的临时副本到 ``output_dir``，资源目录指回原文件目录，使
      ``\\input`` / 图片相对路径仍可解析。任何异常都安全降级为「用原文件」。
    """
    if from_format != "latex":
        return src, None, None, []
    try:
        with open(src, "r", encoding="utf-8") as fh:
            original = fh.read()
    except (OSError, UnicodeDecodeError):
        return src, None, None, []  # 读不了就别动，交原文件给 pandoc（不冒corrupt风险）

    normalized, notes = normalize_latex_for_pandoc(original)
    if normalized == original:
        return src, None, None, []
    try:
        os.makedirs(output_dir, exist_ok=True)
        tmp_path = os.path.join(output_dir, f"{stem}_normalized.tex")
        with open(tmp_path, "w", encoding="utf-8") as fh:
            fh.write(normalized)
    except OSError:
        return src, None, None, []  # 写临时副本失败 → 降级用原文件
    resource_dir = os.path.dirname(os.path.abspath(src)) or "."
    return tmp_path, tmp_path, resource_dir, notes


def _cleanup_temp(path: str | None) -> None:
    """删掉预规整临时副本（存在才删，失败静默——残留临时文件不影响正确性）。"""
    if not path:
        return
    try:
        if os.path.isfile(path):
            os.remove(path)
    except OSError:
        pass


def _remember_source_mutation(src: str, out_path: str):
    """把源文件路径/扩展名与最近产物记入 ws.profile（供后续轮次定位，不动内容）。"""

    def _mutate(ws) -> None:
        ws.profile["source_document_path"] = src
        ws.profile["source_document_ext"] = os.path.splitext(src)[1].lower()
        ws.profile["last_output_path"] = out_path

    return _mutate


def _handle_convert(
    ctx: ToolContext, to_format: str, path: str | None = None,
    two_column: bool = False, three_line_table: bool = True,
) -> str:
    """convert_document 工具入口：调确定性核心，返回人可读文本。"""
    outcome = convert_document_core(ctx, to_format, path, two_column, three_line_table)
    return outcome.message()


def register_convert_document(registry: ToolRegistry, ctx: ToolContext) -> None:
    """把 convert_document 工具注册进 registry。"""
    registry.register(
        name="convert_document",
        description=_DESCRIPTION,
        handler=lambda to_format, path=None, two_column=False, three_line_table=True: (
            _handle_convert(ctx, to_format, path, two_column, three_line_table)
        ),
        parameters=_SCHEMA,
    )


__all__ = ["register_convert_document", "convert_document_core", "ConvertOutcome"]
