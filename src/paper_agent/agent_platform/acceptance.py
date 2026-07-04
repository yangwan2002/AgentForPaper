"""确定性验收层（无 LLM 纯函数核对）。

借鉴编程 agent 用"测试/编译"作客观验证的思路：论文虽无 ground truth，但一批
**可机械检查**的产出信号（乱码、排版是否应用、引用是否闭合、文献数量/年限）可
在任务收尾前做确定性核对，产出结构化 :class:`AcceptanceReport`，据此走"有界自愈
重验"或"诚实上报"（见 acceptance loop / Top_Agent 收尾）。

本模块只提供**纯检查函数 + 数据模型**（Task 1）；验收编排（``AcceptanceChecker`` /
``AcceptanceLoop``）与自愈闭环见 Task 3。各检查为独立函数，便于单测与扩展：

- :func:`detect_mojibake`：文本乱码启发式（U+FFFD / latin-1 高位连续序列 / 异常
  码点比例）——中文经 UTF-8↔GBK/latin-1 错配是最常见的乱码源。
- :func:`check_typesetting_applied`：读回 docx，核对正文段落对齐/行距/首行缩进
  是否等于设定的排版规格。
- :func:`check_citation_closure`：正文 ``[id]`` 集合 vs 已验证文献 id 集合 →
  悬空（正文引用无对应文献）/ 冗余（文献从未被引用）。
- :func:`check_quantity` / :func:`check_recency`：文献数量区间 / 最早年限。

设计取舍：检查结果为 :class:`AcceptanceFinding`（``ok`` + 人类可读 ``detail`` +
``healable`` 是否存在可行修正路径）。内容/工具层问题（排版未应用、悬空引用、占位）
标 ``healable=True``；环境/编码类（缺 pandoc、疑因工具编码导致的乱码）标
``healable=False`` → 直接上报，绝不静默交付、绝不无界重试。
"""

from __future__ import annotations

from dataclasses import dataclass, field
from typing import Callable

from paper_agent.tools.quality_gate import extract_text_citations
from paper_agent.workspace.models import PaperWorkspace


# --------------------------------------------------------------------------- #
# 数据模型
# --------------------------------------------------------------------------- #

@dataclass
class AcceptanceFinding:
    """一条验收发现。

    - ``check``：检查项标识（``"mojibake"`` / ``"typesetting"`` /
      ``"citation_closure"`` / ``"quantity"`` / ``"recency"`` ...）。
    - ``ok``：该项是否通过。
    - ``detail``：人类可读说明（通过 / 未通过原因、证据）。
    - ``healable``：未通过时是否存在可行修正路径（内容/工具层可修 → True；
      环境/编码/代码类 → False，直接上报）。``ok=True`` 时该字段无意义。
    - ``severity``：严重度（默认 ``"high"``）。
    """

    check: str
    ok: bool
    detail: str = ""
    healable: bool = False
    severity: str = "high"


@dataclass
class AcceptanceReport:
    """一次验收的汇总报告。"""

    findings: list[AcceptanceFinding] = field(default_factory=list)

    @property
    def passed(self) -> bool:
        """是否全部检查项通过。"""
        return all(f.ok for f in self.findings)

    @property
    def failures(self) -> list[AcceptanceFinding]:
        """所有未通过项。"""
        return [f for f in self.findings if not f.ok]

    @property
    def healable_failures(self) -> list[AcceptanceFinding]:
        """未通过且存在可行修正路径的项（可触发有界自愈）。"""
        return [f for f in self.findings if not f.ok and f.healable]

    @property
    def blocking_failures(self) -> list[AcceptanceFinding]:
        """未通过且无可行修正路径的项（须诚实上报，不自愈）。"""
        return [f for f in self.findings if not f.ok and not f.healable]

    def to_dict(self) -> dict:
        return {
            "passed": self.passed,
            "findings": [
                {
                    "check": f.check,
                    "ok": f.ok,
                    "detail": f.detail,
                    "healable": f.healable,
                    "severity": f.severity,
                }
                for f in self.findings
            ],
        }


@dataclass
class TaskRequirements:
    """从用户任务解析出的**可测约束**（不臆测：未解析出的约束不纳入验收）。

    各字段默认 ``None``（未指定），语义为「不检查该项」。由 Top_Agent 在任务开始
    时结构化填充（见 Task 4）。

    - ``expected_format``：期望输出格式（``"docx"`` / ``"latex"`` / ``"markdown"``）。
    - ``typesetting``：期望排版规格（:class:`Typesetting` 或等价 dict）。
    - ``reference_count_min`` / ``reference_count_max``：文献数量区间。
    - ``min_year``：文献最早年限。
    - ``require_citation_closure``：是否核对引用闭合（默认 True——闭合是通用正确性）。
    """

    expected_format: str | None = None
    typesetting: object | None = None
    reference_count_min: int | None = None
    reference_count_max: int | None = None
    min_year: int | None = None
    require_citation_closure: bool = True

    def has_any(self) -> bool:
        """是否存在任一可测约束（无则 Top_Agent 不做多余验收，保持既有路径行为）。"""
        return any(
            v is not None
            for v in (
                self.expected_format,
                self.typesetting,
                self.reference_count_min,
                self.reference_count_max,
                self.min_year,
            )
        ) or self.require_citation_closure


# --------------------------------------------------------------------------- #
# 检查 1：乱码检测（启发式，无 LLM）
# --------------------------------------------------------------------------- #

# latin-1 高位区（U+0080..U+00FF）：中文 UTF-8 字节被误按 latin-1 解码时，
# 一个汉字（3 字节）会变成 3 个该区间字符，故**连续**该区间字符是强乱码信号。
_LATIN1_HIGH_LO = 0x80
_LATIN1_HIGH_HI = 0xFF

# 连续 latin-1 高位字符达到此长度即判定为乱码序列（3 起步——一个汉字的字节数）。
_MOJIBAKE_RUN_THRESHOLD = 3

# 异常码点占比阈值：疑似乱码字符占总字符比例超过此值即判乱码。
_MOJIBAKE_RATIO_THRESHOLD = 0.20


def _is_latin1_high(ch: str) -> bool:
    return _LATIN1_HIGH_LO <= ord(ch) <= _LATIN1_HIGH_HI


def _max_latin1_run(text: str) -> int:
    """返回文本中最长的「连续 latin-1 高位字符」游程长度。"""
    longest = 0
    current = 0
    for ch in text:
        if _is_latin1_high(ch):
            current += 1
            longest = max(longest, current)
        else:
            current = 0
    return longest


def detect_mojibake(text: str) -> tuple[bool, str]:
    """检测文本是否疑似乱码。

    返回 ``(是否乱码, 证据说明)``。三条互补启发式，任一触发即判乱码：

    1. **U+FFFD 替换符**：解码失败留下的替换字符，出现即强信号。
    2. **连续 latin-1 高位序列**：中文 UTF-8 字节被误按 latin-1/GBK 解码的典型
       特征——一串 ``Ã¦Â``、``é«``、``â``... 连续高位字符（游程 ≥ 3）。
    3. **异常码点占比**：latin-1 高位字符 + 替换符占总字符比例超阈值。

    对正常中文（CJK 统一表意区，码点远高于 latin-1 区）与正常英文（ASCII）均返回
    ``(False, ...)``。
    """
    text = text or ""
    if not text:
        return False, "空文本"

    replacement_count = text.count("\ufffd")
    if replacement_count > 0:
        return True, f"含 {replacement_count} 个 U+FFFD 替换符（解码失败残留）"

    longest_run = _max_latin1_run(text)
    if longest_run >= _MOJIBAKE_RUN_THRESHOLD:
        return True, (
            f"检出连续 {longest_run} 个 latin-1 高位字符——"
            f"典型的 UTF-8↔GBK/latin-1 编码错配乱码特征"
        )

    high_count = sum(1 for ch in text if _is_latin1_high(ch))
    total = len(text)
    ratio = high_count / total if total else 0.0
    if ratio > _MOJIBAKE_RATIO_THRESHOLD:
        return True, (
            f"异常码点占比 {ratio:.0%} 超过阈值 {_MOJIBAKE_RATIO_THRESHOLD:.0%}"
        )

    return False, "未检出乱码特征"


# --------------------------------------------------------------------------- #
# 检查 2：排版是否应用（读回 docx 核对）
# --------------------------------------------------------------------------- #

# python-docx WD_ALIGN_PARAGRAPH 枚举值 → 语义对齐名。
_ALIGN_ENUM_TO_NAME = {0: "left", 1: "center", 2: "right", 3: "justify"}

# 首行缩进容差（EMU/磅换算的浮点误差；1 磅 ≈ 12700 EMU）。
_INDENT_TOLERANCE_PT = 0.5
# 行距容差（磅）。
_LINE_SPACING_TOLERANCE_PT = 0.5


def _spec_get(spec, key: str):
    """从 :class:`Typesetting` 或 dict 统一取字段值。"""
    if spec is None:
        return None
    if isinstance(spec, dict):
        return spec.get(key)
    return getattr(spec, key, None)


def _parse_first_line_indent_pt(value) -> float | None:
    """把首行缩进规格解析为磅值。

    支持 ``"2ch"``（按常见惯例 1 中文字符 ≈ 半个正文行高，此处只做「非零即已缩进」
    的存在性核对，不追求精确磅值）与纯磅值字符串/数字。无法解析返回 ``None``。
    """
    if value is None:
        return None
    if isinstance(value, (int, float)):
        return float(value)
    text = str(value).strip().lower()
    if text.endswith("ch"):
        # "2ch" 语义为「首行缩进 2 个字符」——此处仅要求缩进存在（>0），
        # 精确磅值依字体字号而定，不在确定性核对范围内。
        try:
            return float(text[:-2]) if float(text[:-2]) else 0.0
        except ValueError:
            return None
    if text.endswith("pt"):
        text = text[:-2]
    try:
        return float(text)
    except ValueError:
        return None


def check_typesetting_applied(docx_path: str, spec) -> AcceptanceFinding:
    """读回 docx，核对正文段落的对齐/行距/首行缩进是否等于设定的排版规格。

    仅核对 ``spec`` 中**已指定**的字段（未指定的不检查）。只要存在足够多的正文
    段落**未**应用某项已指定规格即判未通过（healable=True——可重新应用排版）。

    读 docx 失败（缺 python-docx、文件损坏等）记为环境类未通过
    （healable=False → 上报），不崩溃。
    """
    want_align = _spec_get(spec, "alignment")
    want_spacing = _spec_get(spec, "line_spacing")
    want_indent_raw = _spec_get(spec, "first_line_indent")
    want_indent_pt = _parse_first_line_indent_pt(want_indent_raw)

    if want_align is None and want_spacing is None and want_indent_raw is None:
        return AcceptanceFinding(
            check="typesetting", ok=True, detail="未指定排版规格，跳过核对"
        )

    try:
        import docx  # noqa: WPS433 - 惰性导入可选依赖
    except ImportError:
        return AcceptanceFinding(
            check="typesetting",
            ok=False,
            detail="无法核对排版：未安装 python-docx",
            healable=False,
        )

    try:
        document = docx.Document(docx_path)
    except Exception as exc:  # noqa: BLE001 - 读文件失败记为环境类未通过
        return AcceptanceFinding(
            check="typesetting",
            ok=False,
            detail=f"无法打开 docx 核对排版：{exc}",
            healable=False,
        )

    mismatches: list[str] = []
    body_paras = [p for p in document.paragraphs if (p.text or "").strip()]
    checked = 0
    for para in body_paras:
        # 标题段落（Heading 样式）不受正文排版约束，跳过。
        style_name = getattr(getattr(para, "style", None), "name", "") or ""
        if style_name.lower().startswith("heading") or style_name.startswith("标题"):
            continue
        checked += 1
        fmt = para.paragraph_format

        if want_align is not None:
            actual = _ALIGN_ENUM_TO_NAME.get(
                int(fmt.alignment) if fmt.alignment is not None else -1
            )
            if actual != want_align:
                mismatches.append(f"对齐={actual or '默认'}≠{want_align}")

        if want_spacing is not None and fmt.line_spacing is not None:
            actual_sp = float(fmt.line_spacing)
            # python-docx 对固定磅行距返回 Pt 对象（EMU）；倍数行距返回浮点。
            if hasattr(fmt.line_spacing, "pt"):
                actual_sp = float(fmt.line_spacing.pt)
            if abs(actual_sp - float(want_spacing)) > _LINE_SPACING_TOLERANCE_PT:
                mismatches.append(f"行距={actual_sp}≠{want_spacing}")
        elif want_spacing is not None and fmt.line_spacing is None:
            mismatches.append("行距=未设置")

        if want_indent_pt is not None:
            indent = fmt.first_line_indent
            actual_indent_pt = float(indent.pt) if indent is not None else 0.0
            if want_indent_pt > 0 and actual_indent_pt <= _INDENT_TOLERANCE_PT:
                mismatches.append("首行缩进=未应用")

    if checked == 0:
        return AcceptanceFinding(
            check="typesetting",
            ok=True,
            detail="无正文段落可核对（仅结构性内容）",
        )

    if mismatches:
        # 汇总去重的失配类型（避免逐段落刷屏）。
        unique = sorted(set(mismatches))
        return AcceptanceFinding(
            check="typesetting",
            ok=False,
            detail=(
                f"{len(body_paras)} 个正文段落中检出排版失配："
                + "；".join(unique)
            ),
            healable=True,
        )

    return AcceptanceFinding(
        check="typesetting", ok=True, detail=f"{checked} 个正文段落排版规格已应用"
    )


# --------------------------------------------------------------------------- #
# 检查 3：引用闭合
# --------------------------------------------------------------------------- #

def cited_reference_ids(ws: PaperWorkspace) -> set[str]:
    """扫描各章节正文与记录字段，返回被正文实际引用的文献 id 集合。

    合并两条互补路径：正文中出现的 ``[id]`` 标注（``extract_text_citations``）与
    章节记录的 ``cited_reference_ids``——避免「记录未同步」或「正文标注未记录」
    任一遗漏。
    """
    cited: set[str] = set()
    for draft in ws.section_drafts.values():
        content = draft.content or ""
        cited.update(extract_text_citations(content))
        cited.update(draft.cited_reference_ids)
    return cited


def check_citation_closure(ws: PaperWorkspace) -> AcceptanceFinding:
    """核对引用闭合：悬空引用（正文引用无对应文献）与冗余文献（从未被引用）。

    - **悬空**：正文 ``[id]`` 引用的 id 不在已验证文献库 → 未通过（healable：
      可补检索或删标注）。
    - **冗余**：已验证文献从未被任何章节引用 → 记入 detail 提示（不单独判失败，
      因导出层已做引用闭合只列被引用者；此处作为可观测提示）。

    无已验证文献且无正文引用 → 通过（无可核对内容）。
    """
    verified = ws.verified_reference_ids()
    cited = cited_reference_ids(ws)

    dangling = sorted(cited - verified)
    redundant = sorted(verified - cited)

    if dangling:
        detail = f"检出 {len(dangling)} 处悬空引用（正文引用了未核验文献）：{dangling}"
        if redundant:
            detail += f"；另有 {len(redundant)} 篇已验证文献未被引用"
        return AcceptanceFinding(
            check="citation_closure", ok=False, detail=detail, healable=True,
        )

    detail = "引用闭合：正文引用均有对应文献"
    if redundant:
        detail += f"（{len(redundant)} 篇已验证文献未被引用，导出时将不列入参考文献表）"
    return AcceptanceFinding(check="citation_closure", ok=True, detail=detail)


# --------------------------------------------------------------------------- #
# 检查 4：文献数量与年限
# --------------------------------------------------------------------------- #

def check_quantity(
    ws: PaperWorkspace, lo: int | None = None, hi: int | None = None
) -> AcceptanceFinding:
    """核对已验证文献数量是否落在 ``[lo, hi]`` 区间内。

    ``lo`` / ``hi`` 任一为 ``None`` 表示该侧不设限。二者均为 ``None`` → 跳过。
    """
    if lo is None and hi is None:
        return AcceptanceFinding(
            check="quantity", ok=True, detail="未指定文献数量约束，跳过核对"
        )
    count = len(ws.verified_reference_ids())
    if lo is not None and count < lo:
        return AcceptanceFinding(
            check="quantity",
            ok=False,
            detail=f"已验证文献 {count} 篇，少于要求下限 {lo} 篇",
            healable=True,
        )
    if hi is not None and count > hi:
        return AcceptanceFinding(
            check="quantity",
            ok=False,
            detail=f"已验证文献 {count} 篇，超过要求上限 {hi} 篇",
            healable=True,
        )
    return AcceptanceFinding(
        check="quantity", ok=True, detail=f"已验证文献 {count} 篇，满足数量约束"
    )


def check_recency(ws: PaperWorkspace, min_year: int | None = None) -> AcceptanceFinding:
    """核对是否存在早于 ``min_year`` 的已验证文献（年限约束）。

    仅统计有年份且被引用/入库的已验证文献；``min_year`` 为 ``None`` → 跳过。
    年份缺失（``None``）的文献不计入违规（无法判定）。
    """
    if min_year is None:
        return AcceptanceFinding(
            check="recency", ok=True, detail="未指定年限约束，跳过核对"
        )
    stale = [
        r
        for r in ws.verified_references
        if r.verified and r.year is not None and r.year < min_year
    ]
    if stale:
        years = sorted({r.year for r in stale})
        return AcceptanceFinding(
            check="recency",
            ok=False,
            detail=f"检出 {len(stale)} 篇早于 {min_year} 年的文献（年份：{years}）",
            healable=True,
        )
    return AcceptanceFinding(
        check="recency", ok=True, detail=f"所有文献均不早于 {min_year} 年"
    )


# --------------------------------------------------------------------------- #
# 验收编排：AcceptanceChecker（组合各检查）
# --------------------------------------------------------------------------- #

# 期望输出格式 → 产物文件扩展名。
_FORMAT_EXT = {"docx": ".docx", "latex": ".tex", "markdown": ".md"}


def _read_export_text(path: str) -> str | None:
    """读回导出文件的可比对文本；无法读取返回 ``None``（跳过该文件的乱码核对）。

    - ``.docx``：用 python-docx 提取段落文本（若内部存的是乱码字符，此处如实返回，
      供 :func:`detect_mojibake` 检出）。
    - 其余（``.md`` / ``.tex`` / 纯文本）：以 UTF-8 读取，``errors="replace"`` 使
      非法字节变成 U+FFFD —— 恰是乱码的强信号。
    """
    low = path.lower()
    try:
        if low.endswith(".docx"):
            import docx  # noqa: WPS433 - 惰性导入可选依赖

            document = docx.Document(path)
            return "\n".join(p.text for p in document.paragraphs)
        with open(path, encoding="utf-8", errors="replace") as fh:
            return fh.read()
    except Exception:  # noqa: BLE001 - 读文件失败即跳过（不误报乱码）
        return None


class AcceptanceChecker:
    """据任务的可测需求，对工作区 + 导出产物做确定性核对（无 LLM），产出报告。

    组合本模块的各纯检查函数，只核对 ``requirements`` 中**已指定**的约束（未指定
    的不臆测）。乱码为环境/编码类（``healable=False``，诚实上报）；引用闭合/数量/
    年限/排版为内容/工具层（``healable=True``，可自愈）。
    """

    def check(
        self,
        ws: PaperWorkspace,
        export_files: list[str],
        requirements: TaskRequirements,
    ) -> AcceptanceReport:
        findings: list[AcceptanceFinding] = []
        files = list(export_files or [])

        # 期望格式是否产出对应文件。
        if requirements.expected_format:
            ext = _FORMAT_EXT.get(requirements.expected_format)
            if ext and not any(f.lower().endswith(ext) for f in files):
                findings.append(
                    AcceptanceFinding(
                        check="format",
                        ok=False,
                        detail=f"未产出期望格式（{requirements.expected_format}）的文件",
                        healable=True,
                    )
                )

        if requirements.require_citation_closure:
            findings.append(check_citation_closure(ws))

        if (
            requirements.reference_count_min is not None
            or requirements.reference_count_max is not None
        ):
            findings.append(
                check_quantity(
                    ws,
                    requirements.reference_count_min,
                    requirements.reference_count_max,
                )
            )

        if requirements.min_year is not None:
            findings.append(check_recency(ws, requirements.min_year))

        # 逐导出文件核对：乱码（所有格式）+ 排版（仅 docx，且指定了排版规格）。
        for path in files:
            text = _read_export_text(path)
            if text is not None:
                is_bad, evidence = detect_mojibake(text)
                findings.append(
                    AcceptanceFinding(
                        check="mojibake",
                        ok=not is_bad,
                        detail=f"{path}：{evidence}",
                        healable=False,  # 编码/工具类，诚实上报，不静默重试
                    )
                )
            if requirements.typesetting is not None and path.lower().endswith(".docx"):
                findings.append(
                    check_typesetting_applied(path, requirements.typesetting)
                )

        return AcceptanceReport(findings=findings)


# --------------------------------------------------------------------------- #
# 诚实上报 + 有界自愈闭环：AcceptanceLoop
# --------------------------------------------------------------------------- #

@dataclass
class DeliveryOutcome:
    """一次"导出→验收→(自愈/上报)→交付"闭环的结论（面向用户诚实反馈）。

    - ``delivered``：最终产出是否全部通过验收。
    - ``report``：最后一轮验收报告。
    - ``healed``：已自愈的检查项标识列表。
    - ``unresolved``：未解决项的人类可读说明（诚实上报，绝不静默交付）。
    - ``export_files``：最终产出文件路径。
    - ``heal_rounds``：实际自愈轮数（≤ max_heal_rounds，证明有界终止）。
    """

    delivered: bool
    report: AcceptanceReport
    healed: list[str] = field(default_factory=list)
    unresolved: list[str] = field(default_factory=list)
    export_files: list[str] = field(default_factory=list)
    heal_rounds: int = 0


class AcceptanceLoop:
    """有界自愈闭环编排（无 LLM 逻辑本身；自愈动作经注入的回调完成）。

    依赖注入（依赖倒置，便于测试与解耦）：
    - ``checker``：:class:`AcceptanceChecker`。
    - ``export_fn(ws) -> list[str]``：导出并返回产出文件路径。
    - ``heal_fn(session, findings) -> None``（可选）：针对可自愈发现做修正——其内部
      改工作区仍须经既有 ``commit``/护栏/单一写路径（本类不直接写工作区）。为 ``None``
      时不自愈（只导出→验收→上报）。

    流程：导出 → 验收 → 若有可自愈失败且未超 ``max_heal_rounds`` → 调 ``heal_fn`` 修
    正 → 重导出重验；否则收敛。无可行修正路径的失败（乱码/环境类）不触发自愈，直接
    进入 ``unresolved`` 诚实上报。任意情况下有限步终止（Property 4）。
    """

    def __init__(
        self,
        checker: AcceptanceChecker,
        export_fn: Callable[[PaperWorkspace], list[str]],
        heal_fn: Callable[["object", list[AcceptanceFinding]], None] | None = None,
    ) -> None:
        self._checker = checker
        self._export_fn = export_fn
        self._heal_fn = heal_fn

    def run(self, session, requirements: TaskRequirements, *, max_heal_rounds: int = 2) -> DeliveryOutcome:
        ws = session.workspace
        export_files = list(self._export_fn(ws) or [])
        report = self._checker.check(ws, export_files, requirements)

        healed: list[str] = []
        rounds = 0
        while (
            not report.passed
            and report.healable_failures
            and self._heal_fn is not None
            and rounds < max_heal_rounds
        ):
            failures = report.healable_failures
            self._heal_fn(session, failures)
            healed.extend(f.check for f in failures)
            rounds += 1
            export_files = list(self._export_fn(ws) or [])
            report = self._checker.check(ws, export_files, requirements)

        unresolved = [f.detail for f in report.failures]
        return DeliveryOutcome(
            delivered=report.passed,
            report=report,
            healed=healed,
            unresolved=unresolved,
            export_files=export_files,
            heal_rounds=rounds,
        )


__all__ = [
    "AcceptanceFinding",
    "AcceptanceReport",
    "TaskRequirements",
    "DeliveryOutcome",
    "AcceptanceChecker",
    "AcceptanceLoop",
    "detect_mojibake",
    "check_typesetting_applied",
    "cited_reference_ids",
    "check_citation_closure",
    "check_quantity",
    "check_recency",
]
