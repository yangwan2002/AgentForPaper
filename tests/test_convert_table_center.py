"""三线表居中修复测试：docx 表格后处理把表格整体居中（学术惯例）。"""

from __future__ import annotations

import pytest


def test_fix_table_widths_centers_table(tmp_path):
    docx = pytest.importorskip("docx")
    from docx.enum.table import WD_TABLE_ALIGNMENT

    from paper_agent.agent_platform.tools.convert_tool import _fix_table_widths

    path = tmp_path / "t.docx"
    document = docx.Document()
    table = document.add_table(rows=2, cols=3)
    table.cell(0, 0).text = "A"
    table.alignment = WD_TABLE_ALIGNMENT.LEFT  # 先设成靠左
    document.save(str(path))

    _fix_table_widths(str(path))

    reopened = docx.Document(str(path))
    assert reopened.tables[0].alignment == WD_TABLE_ALIGNMENT.CENTER
