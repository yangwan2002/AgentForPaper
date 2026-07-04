"""Hypothesis property tests for venue-templates-figures-tables 图嵌入相关属性。

覆盖设计文档 "Correctness Properties" 中的 Property 10–14（LaTeX / docx 图嵌入、
路径穿越防御）。每条属性一个测试，最少 100 次迭代。

设计约定复现要点：
- LaTeX 图块由 ``LatexExporter._render_tex`` 产出：能定位导出目录内且真实存在的
  ``Figure_Asset`` 时，在 ``figure`` 环境内先写 ``\\includegraphics`` 再写
  ``\\caption`` / ``\\label``；否则仅图题+标签回退并发 ``DEGRADATION`` 事件。
- 路径穿越防御由 ``safe_relative_asset`` 唯一入口负责。
- docx 图块由 ``DocxExporter._render_figure`` 产出：有资产 ``add_picture`` + 图题段落，
  无资产回落 ``figure_id: caption`` 段落并发 ``DEGRADATION``。
"""

from __future__ import annotations

import base64
import os
import re
import tempfile

import pytest
from hypothesis import HealthCheck, given, settings
from hypothesis import strategies as st

from paper_agent.export.asset_paths import safe_relative_asset
from paper_agent.export.docx import DocxExporter
from paper_agent.export.latex import LatexExporter, _escape
from paper_agent.observability.events import Event, EventKind
from paper_agent.workspace.models import (
    FigureRecord,
    InputMode,
    OutputFormat,
    PaperWorkspace,
)

# 已知可用的最小 1x1 PNG（供 LaTeX 路径存在性校验与 docx add_picture 使用）。
_PNG_1x1 = base64.b64decode(
    b"iVBORw0KGgoAAAANSUhEUgAAAAEAAAABCAYAAAAfFcSJAAAADUlEQVR4"
    b"2mP8z8BQDwAEhQGAhKmMIQAAAABJRU5ErkJggg=="
)

# 图题字符：排除控制/代理字符（含换行）以便按行/子串断言，且保证 docx XML 兼容。
_CAPTION_TEXT = st.text(
    alphabet=st.characters(blacklist_categories=("Cc", "Cs")),
    max_size=30,
)


class _ListSink:
    """捕获事件的简单 sink，用于断言降级事件。"""

    def __init__(self) -> None:
        self.events: list[Event] = []

    def emit(self, event: Event) -> None:
        self.events.append(event)


def _base_ws(fmt: OutputFormat) -> PaperWorkspace:
    return PaperWorkspace(
        workspace_id="paper_props",
        input_mode=InputMode.GENERATION,
        output_format=fmt,
        topic_background="x",
    )


def _figure_blocks(tex: str) -> list[str]:
    """抽取 .tex 中每个 figure 环境的内容片段。"""
    return re.findall(r"\\begin\{figure\}.*?\\end\{figure\}", tex, flags=re.DOTALL)


# --------------------------------------------------------------------------- #
# Property 10: LaTeX 图嵌入顺序与一一对应
# --------------------------------------------------------------------------- #
# Feature: venue-templates-figures-tables, Property 10: LaTeX 图嵌入顺序与一一对应
@given(n=st.integers(min_value=1, max_value=5), captions=st.lists(_CAPTION_TEXT, min_size=5, max_size=5))
@settings(max_examples=100, deadline=None, suppress_health_check=[HealthCheck.function_scoped_fixture])
def test_property_10_latex_embed_order_and_one_to_one(n, captions):
    with tempfile.TemporaryDirectory() as out_dir:
        figs = []
        for i in range(n):
            data_ref = f"asset_{i}.png"
            with open(os.path.join(out_dir, data_ref), "wb") as fh:
                fh.write(_PNG_1x1)
            figs.append(
                FigureRecord(figure_id=f"fig{i}", data_ref=data_ref, caption=captions[i])
            )
        ws = _base_ws(OutputFormat.LATEX)
        ws.figures = figs

        LatexExporter().export(ws, out_dir)
        tex = open(os.path.join(out_dir, "paper_props.tex"), encoding="utf-8").read()

        expected = {f.figure_id: f.data_ref for f in figs}
        blocks = _figure_blocks(tex)
        assert len(blocks) == n

        for block in blocks:
            inc = block.find(r"\includegraphics")
            cap = block.find(r"\caption")
            lbl = block.find(r"\label")
            # includegraphics 必须出现且严格早于 caption 与 label。
            assert inc != -1 and cap != -1 and lbl != -1
            assert inc < cap
            assert inc < lbl
            # 每个 figure 环境引用其自身资产（无错配）。
            path = re.search(r"\\includegraphics\{([^}]*)\}", block).group(1)
            fid = re.search(r"\\label\{([^}]*)\}", block).group(1)
            assert path == expected[fid]


# --------------------------------------------------------------------------- #
# Property 11: LaTeX 图路径一致、存在且保留图题/标签
# --------------------------------------------------------------------------- #
# Feature: venue-templates-figures-tables, Property 11: LaTeX 图路径一致、存在且保留图题/标签
@given(n=st.integers(min_value=1, max_value=5), captions=st.lists(_CAPTION_TEXT, min_size=5, max_size=5))
@settings(max_examples=100, deadline=None, suppress_health_check=[HealthCheck.function_scoped_fixture])
def test_property_11_latex_path_consistent_exists_caption_label(n, captions):
    with tempfile.TemporaryDirectory() as out_dir:
        figs = []
        for i in range(n):
            data_ref = f"asset_{i}.png"
            with open(os.path.join(out_dir, data_ref), "wb") as fh:
                fh.write(_PNG_1x1)
            figs.append(
                FigureRecord(figure_id=f"fig{i}", data_ref=data_ref, caption=captions[i])
            )
        ws = _base_ws(OutputFormat.LATEX)
        ws.figures = figs

        result = LatexExporter().export(ws, out_dir)
        tex = open(os.path.join(out_dir, "paper_props.tex"), encoding="utf-8").read()

        for fig in figs:
            rel = safe_relative_asset(out_dir, fig.data_ref)
            assert rel is not None
            # \includegraphics 路径等于落盘相对资产路径。
            assert rf"\includegraphics{{{rel}}}" in tex
            # 文件存在于文件系统。
            abs_path = os.path.abspath(os.path.join(out_dir, rel))
            assert os.path.exists(abs_path)
            # 资产出现在 ExportResult.files。
            assert abs_path in result.files
            # 图题与标签仍保留（图题按 _escape 转义）。
            assert rf"\caption{{{_escape(fig.caption)}}}" in tex
            assert rf"\label{{{fig.figure_id}}}" in tex


# --------------------------------------------------------------------------- #
# Property 12: 缺资产图的 LaTeX 回退
# --------------------------------------------------------------------------- #
# Feature: venue-templates-figures-tables, Property 12: 缺资产图的 LaTeX 回退
@given(n=st.integers(min_value=1, max_value=5), captions=st.lists(_CAPTION_TEXT, min_size=5, max_size=5))
@settings(max_examples=100, deadline=None, suppress_health_check=[HealthCheck.function_scoped_fixture])
def test_property_12_latex_missing_asset_fallback(n, captions):
    with tempfile.TemporaryDirectory() as out_dir:
        # data_ref 指向不落盘的文件：路径安全但资产不存在 → 缺资产回退。
        figs = [
            FigureRecord(
                figure_id=f"fig{i}", data_ref=f"missing_{i}.png", caption=captions[i]
            )
            for i in range(n)
        ]
        ws = _base_ws(OutputFormat.LATEX)
        ws.figures = figs
        sink = _ListSink()

        LatexExporter(sink=sink).export(ws, out_dir)
        tex = open(os.path.join(out_dir, "paper_props.tex"), encoding="utf-8").read()

        blocks = _figure_blocks(tex)
        assert len(blocks) == n
        for block, fig in zip(blocks, figs):
            # 不生成 \includegraphics。
            assert r"\includegraphics" not in block
            # 仍保留图题与标签。
            assert rf"\caption{{{_escape(fig.caption)}}}" in block
            assert rf"\label{{{fig.figure_id}}}" in block

        # 每个缺资产图记录一条 figure_embed_latex 的降级事件。
        degradations = [
            e
            for e in sink.events
            if e.kind is EventKind.DEGRADATION
            and e.data.get("feature") == "figure_embed_latex"
        ]
        assert len(degradations) == n


# --------------------------------------------------------------------------- #
# Property 13: 图像路径穿越防御
# --------------------------------------------------------------------------- #
def _candidate_paths():
    name = st.text(
        alphabet=st.characters(whitelist_categories=("Ll", "Lu", "Nd")),
        min_size=1,
        max_size=8,
    )
    inside = name.map(lambda s: f"{s}.png")
    inside_sub = st.tuples(name, name).map(lambda t: f"{t[0]}/{t[1]}.png")
    traversal = st.tuples(st.integers(1, 4), name).map(
        lambda t: "/".join([".."] * t[0]) + f"/{t[1]}.png"
    )
    abs_outside = name.map(
        lambda s: os.path.abspath(os.path.join(tempfile.gettempdir(), f"outside_{s}.png"))
    )
    return st.one_of(inside, inside_sub, traversal, abs_outside)


# Feature: venue-templates-figures-tables, Property 13: 图像路径穿越防御
@given(candidate=_candidate_paths())
@settings(max_examples=100, suppress_health_check=[HealthCheck.function_scoped_fixture])
def test_property_13_asset_path_traversal_defense(candidate):
    with tempfile.TemporaryDirectory() as out_dir:
        base = os.path.realpath(os.path.abspath(out_dir))
        rel = safe_relative_asset(out_dir, candidate)
        if rel is None:
            return
        # 非 None 时必须是位于 out_dir 之内的相对路径。
        assert not os.path.isabs(rel)
        resolved = os.path.realpath(os.path.join(base, rel))
        assert resolved == base or os.path.commonpath([base, resolved]) == base


# --------------------------------------------------------------------------- #
# Property 14: docx 图嵌入与图题一一对应
# --------------------------------------------------------------------------- #
# Feature: venue-templates-figures-tables, Property 14: docx 图嵌入与图题一一对应
@given(
    n_asset=st.integers(min_value=1, max_value=4),
    n_missing=st.integers(min_value=0, max_value=3),
    captions=st.lists(_CAPTION_TEXT, min_size=7, max_size=7),
)
@settings(max_examples=100, deadline=None, suppress_health_check=[HealthCheck.function_scoped_fixture])
def test_property_14_docx_embed_one_to_one(n_asset, n_missing, captions):
    pytest.importorskip("docx")

    with tempfile.TemporaryDirectory() as out_dir:
        asset_figs = []
        for i in range(n_asset):
            data_ref = f"asset_{i}.png"
            with open(os.path.join(out_dir, data_ref), "wb") as fh:
                fh.write(_PNG_1x1)
            asset_figs.append(
                FigureRecord(figure_id=f"a{i}", data_ref=data_ref, caption=captions[i])
            )
        missing_figs = [
            FigureRecord(
                figure_id=f"m{i}", data_ref=f"missing_{i}.png", caption=captions[n_asset + i]
            )
            for i in range(n_missing)
        ]
        ws = _base_ws(OutputFormat.DOCX)
        ws.figures = asset_figs + missing_figs
        sink = _ListSink()

        result = DocxExporter(sink=sink).export(ws, out_dir)

        import docx

        document = docx.Document(result.files[0])

        # 内联图片数量 == 有资产图数量。
        assert len(document.inline_shapes) == n_asset

        para_texts = [p.text for p in document.paragraphs]
        # 每张内联图片下方输出其对应图题文本。
        for fig in asset_figs:
            assert f"{fig.figure_id}: {fig.caption}" in para_texts
        # 无资产图回落文字段落 + 降级事件。
        for fig in missing_figs:
            assert f"{fig.figure_id}: {fig.caption}" in para_texts

        docx_degradations = [
            e
            for e in sink.events
            if e.kind is EventKind.DEGRADATION
            and e.data.get("feature") == "figure_embed_docx"
        ]
        assert len(docx_degradations) == n_missing
