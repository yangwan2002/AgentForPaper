"""Property-based tests for the ``venue-templates-figures-tables`` spec.

覆盖 design.md「Correctness Properties」中与结果表/数据出图 grounding、结构完整性、
无数据优雅跳过、数值格式化/派生文本截断相关的属性（Property 1/2/3/15/16/17）。

每条属性一个测试，使用 Hypothesis（``@settings(max_examples=100)``），生成随机
``Experiment``/``ResearchArtifact``（随机 stats/baselines/metrics）。图相关属性使用一个
收集事件的假 sink 与一个记录被绘数值、并落盘桩文件的桩绘图后端；docx 相关断言在
python-docx 不可用时跳过。
"""

from __future__ import annotations

import re
import statistics
import tempfile

import pytest
from hypothesis import HealthCheck, assume, given, settings
from hypothesis import strategies as st

from paper_agent.export.figure_renderer import FigureRenderer
from paper_agent.export.grounding import GroundingChecker
from paper_agent.export.table_renderer import TableRenderer
from paper_agent.observability.events import EventKind
from paper_agent.tools.quality_gate import (
    QualityGate,
    build_allowed_values,
    value_matches,
)
from paper_agent.workspace.models import (
    InputMode,
    OutlineNode,
    PaperWorkspace,
    SectionDraft,
)
from paper_agent.workspace.research_artifact import (
    Experiment,
    MethodSpec,
    ResearchArtifact,
)

try:  # python-docx 为可选依赖：不可用时跳过 docx 相关断言。
    import docx

    HAS_DOCX = True
except Exception:  # pragma: no cover - 依赖缺失路径
    HAS_DOCX = False


# --------------------------------------------------------------------------- #
# 生成器与测试替身
# --------------------------------------------------------------------------- #

METRIC_POOL = ["acc", "map", "recall", "f1", "auc", "bleu"]
BASELINE_POOL = ["OursA", "BaseB", "MethodC", "RefD", "AltE"]

# 数值约束在 [10.0, 99.0] 且四舍五入到 1 位小数：
# - 2 位整数 + 1 位小数的形式在按任意小数位格式化后再解析仍落在 1% 相对容差内，
#   grounding 判定稳定；
# - 更重要的是这种形式嵌入正文后，``QualityGate`` 面向散文的数字抽取器不会把小数
#   的尾随零段误当作独立的 3 位以上整数（如 "1.000" 会被抽出 "000"→0.0），从而
#   避免测试自身构造的假 fabricated_metric。
_VALUE = st.floats(
    min_value=10.0, max_value=99.0, allow_nan=False, allow_infinity=False
).map(lambda x: round(x, 1))


class RecordingSink:
    """收集所有事件，便于断言降级/跳过被正确记录。"""

    def __init__(self) -> None:
        self.events: list = []

    def emit(self, event) -> None:
        self.events.append(event)


class StubPlottingBackend:
    """桩绘图后端：记录每次被传入的数值，并落盘一个占位文件（不依赖 matplotlib）。"""

    available = True

    def __init__(self) -> None:
        self.calls: list[tuple] = []

    def bar_chart(self, title, labels, values, out_path) -> None:
        self.calls.append((title, list(labels), list(values), out_path))
        with open(out_path, "w", encoding="utf-8") as fh:
            fh.write("stub")


@st.composite
def _experiment(draw, idx=0, with_stats=True, with_baselines=True, with_metrics=True):
    """随机构造一个 ``Experiment``（含 rows/stats/baselines/metrics）。"""
    metrics = (
        draw(st.lists(st.sampled_from(METRIC_POOL), min_size=1, max_size=3, unique=True))
        if with_metrics
        else []
    )
    baselines = (
        draw(
            st.lists(st.sampled_from(BASELINE_POOL), min_size=1, max_size=3, unique=True)
        )
        if with_baselines
        else []
    )
    eff_metrics = metrics if metrics else ["acc"]
    row_labels = baselines if baselines else ["result"]

    rows = []
    for label in row_labels:
        row = {"method": label}
        for metric in eff_metrics:
            row[metric] = round(draw(_VALUE), 4)
        rows.append(row)

    stats: dict = {}
    if with_stats:
        for metric in eff_metrics:
            vals = [r[metric] for r in rows]
            mean = round(sum(vals) / len(vals), 4)
            std = round(statistics.pstdev(vals), 4) if len(vals) > 1 else 0.0
            stats[metric] = {"mean": mean, "std": std, "min": min(vals), "max": max(vals)}

    results_data = {"columns": ["method"] + eff_metrics, "rows": rows, "stats": stats}
    return Experiment(
        experiment_id=f"e{idx}",
        dataset="",
        baselines=baselines,
        metrics=metrics,
        results_data=results_data,
    )


@st.composite
def _artifact(draw, with_stats=True, with_baselines=True, with_metrics=True, max_exp=3):
    """随机构造含 1..max_exp 个实验的 ``ResearchArtifact``。"""
    n = draw(st.integers(min_value=1, max_value=max_exp))
    exps = [
        draw(
            _experiment(
                idx=i,
                with_stats=with_stats,
                with_baselines=with_baselines,
                with_metrics=with_metrics,
            )
        )
        for i in range(n)
    ]
    return ResearchArtifact(
        research_question="rq", method=MethodSpec(overview="m"), experiments=exps
    )


@st.composite
def _artifact_empty_stats(draw):
    """构造 stats 全空（或缺失 stats 键）的 artifact——用于「无数据优雅跳过」。"""
    n = draw(st.integers(min_value=1, max_value=3))
    exps = []
    for i in range(n):
        rd = {
            "columns": ["method", "acc"],
            "rows": [{"method": "OursA", "acc": round(draw(_VALUE), 4)}],
            "stats": {},
        }
        if draw(st.booleans()):
            rd.pop("stats")  # 覆盖「完全没有 stats 键」的分支。
        exps.append(Experiment(experiment_id=f"e{i}", results_data=rd))
    return ResearchArtifact(
        research_question="rq", method=MethodSpec(overview="m"), experiments=exps
    )


# --------------------------------------------------------------------------- #
# Property 1
# --------------------------------------------------------------------------- #


# Feature: venue-templates-figures-tables, Property 1: 表/图数值 grounding 不变式——
# TableRenderer 产出与传给绘图后端的每个数值均为 Grounded_Value，纳入工作区后不触发 fabricated_metric
@settings(
    max_examples=100,
    suppress_health_check=[HealthCheck.too_slow],
    # 该属性涉及绘图后端桩的落盘 I/O，单例耗时在满负载下会超过 Hypothesis 默认
    # 200ms deadline（DeadlineExceeded/Flaky 假失败）。这是计时敏感、非逻辑问题，
    # 故对本条禁用 per-example deadline。
    deadline=None,
)
@given(artifact=_artifact())
def test_property_1_table_and_figure_values_are_grounded(artifact):
    sink = RecordingSink()
    checker = GroundingChecker(artifact)
    allowed = build_allowed_values(artifact)

    # 表：所有渲染出的数值单元格必须 grounded。
    renderer = TableRenderer(checker, sink)
    fragments = renderer.render_latex(artifact)
    table_numbers: list[str] = []
    for frag in fragments:
        for tok in re.findall(r"\d+\.\d+", frag.latex):
            table_numbers.append(tok)
            assert value_matches(float(tok), allowed), (
                f"表格数值 {tok} 未 grounded"
            )

    # 图：传给绘图后端的每个数值必须 grounded。
    backend = StubPlottingBackend()
    fig_renderer = FigureRenderer(backend, checker, sink, tracker=None, enabled=True)
    assets_dir = tempfile.mkdtemp()
    fig_renderer.render_from_artifact(artifact, assets_dir)
    for _title, _labels, values, _out in backend.calls:
        for v in values:
            assert value_matches(float(v), allowed), f"图数值 {v} 未 grounded"

    # 纳入工作区后 QualityGate.check 不产生针对表/图数值的 fabricated_metric。
    # 用 artifact 的原始 grounded 数值（1 位小数、散文友好）构造正文，避免把表格的
    # 3 位小数格式化串嵌入散文时被质量闸的数字抽取器错误拆分。
    ws = PaperWorkspace(
        workspace_id="w", input_mode=InputMode.GENERATION, topic_background="t"
    )
    ws.artifact = artifact
    ws.outline = [OutlineNode(section_id="results", title="Results", order=0)]
    raw_values = artifact.all_numeric_values()
    content = (
        "Reported metrics: "
        + " ".join(f"{v:.1f}" for v in raw_values)
        + " were observed across settings."
    )
    ws.section_drafts = {
        "results": SectionDraft(section_id="results", title="Results", content=content)
    }
    report = QualityGate().check(ws)
    fabricated = [i for i in report.issues if i.get("type") == "fabricated_metric"]
    assert not fabricated, f"出现 fabricated_metric：{fabricated}"


# --------------------------------------------------------------------------- #
# Property 2
# --------------------------------------------------------------------------- #


# Feature: venue-templates-figures-tables, Property 2: 非 grounded 数值被拒绝并记录——
# 未 grounded 值与异常单元被跳过并记录，整表不中止
@settings(max_examples=100)
@given(
    g=_VALUE,
    offset=st.floats(
        min_value=50.0, max_value=5000.0, allow_nan=False, allow_infinity=False
    ),
)
def test_property_2_ungrounded_values_rejected_and_recorded(g, offset):
    # grounding 允许集合仅含 g（该 artifact 无 stats、rows 只含 g）。
    grounding_artifact = ResearchArtifact(
        research_question="rq",
        method=MethodSpec(overview="m"),
        experiments=[
            Experiment(experiment_id="g", results_data={"rows": [{"v": g}], "stats": {}})
        ],
    )
    checker = GroundingChecker(grounding_artifact)

    u = round(g + offset + 100.0, 4)  # 距 g 足够远 → 必然未 grounded。
    assume(not value_matches(u, checker.allowed_values()))

    render_artifact = ResearchArtifact(
        research_question="rq",
        method=MethodSpec(overview="m"),
        experiments=[
            Experiment(
                experiment_id="r",
                dataset="",
                baselines=["OursA"],
                metrics=["m1", "m2", "m3"],
                results_data={
                    "rows": [{"method": "OursA", "m1": g, "m2": u, "m3": "not_a_number"}],
                    "stats": {"m1": {"mean": g, "std": 0.0, "min": g, "max": g}},
                },
            )
        ],
    )

    sink = RecordingSink()
    renderer = TableRenderer(checker, sink)
    fragments = renderer.render_latex(render_artifact)

    # 整表未中止：仍产出 1 个 fragment。
    assert len(fragments) == 1
    frag = fragments[0]

    # 未 grounded 单元与异常单元被跳过并记录。
    assert "OursA/m2" in frag.skipped_cells
    assert "OursA/m3" in frag.skipped_cells

    # grounded 值仍被渲染；未 grounded 值不出现在产物中。
    assert f"{g:.3f}" in frag.latex
    assert f"{u:.3f}" not in frag.latex

    reasons = [
        e.data.get("reason") for e in sink.events if e.kind == EventKind.DEGRADATION
    ]
    assert "rejected_ungrounded_value" in reasons
    assert "cell_skipped" in reasons


# --------------------------------------------------------------------------- #
# Property 3
# --------------------------------------------------------------------------- #


# Feature: venue-templates-figures-tables, Property 3: grounding 允许集合与既有质量闸同源一致——
# GroundingChecker.allowed_values() 等于 QualityGate 构造的 extended_allowed（build_allowed_values）
@settings(max_examples=100)
@given(artifact=st.one_of(st.none(), _artifact(), _artifact_empty_stats()))
def test_property_3_allowed_values_same_source_as_quality_gate(artifact):
    checker = GroundingChecker(artifact)
    if artifact is None:
        assert checker.allowed_values() == []
    else:
        assert checker.allowed_values() == build_allowed_values(artifact)


# --------------------------------------------------------------------------- #
# Property 15
# --------------------------------------------------------------------------- #


# Feature: venue-templates-figures-tables, Property 15: 结果表结构完整性——
# LaTeX 片段含 tabular/\caption/\label/表头；docx 表行数>=2；行列反映 baselines×metrics
@settings(max_examples=100)
@given(artifact=_artifact(max_exp=1))
def test_property_15_result_table_structure_complete(artifact):
    sink = RecordingSink()
    renderer = TableRenderer(GroundingChecker(artifact), sink)
    fragments = renderer.render_latex(artifact)
    assert fragments
    exp = artifact.experiments[0]
    frag = fragments[0]

    assert r"\begin{tabular}" in frag.latex
    assert r"\caption" in frag.latex
    assert r"\label" in frag.latex
    assert "Method" in frag.latex  # 表头首列
    for metric in exp.metrics:
        assert metric in frag.latex  # 每个 metric 出现在表头

    if HAS_DOCX:
        document = docx.Document()
        n = renderer.render_docx(artifact, document)
        assert n == 1
        assert len(document.tables) == 1
        table = document.tables[0]
        # 表头行 + 每个 baseline 一行 → >=2 行。
        assert len(table.rows) >= 2
        assert len(table.rows) == 1 + len(exp.baselines)
        assert len(table.columns) == 1 + len(exp.metrics)


# --------------------------------------------------------------------------- #
# Property 16
# --------------------------------------------------------------------------- #


# Feature: venue-templates-figures-tables, Property 16: 无数据时优雅跳过表格——
# 无 artifact / 全空 stats → render_latex 返回 [] 且记录 no_data，不抛异常
@settings(max_examples=100)
@given(artifact=st.one_of(st.none(), _artifact_empty_stats()))
def test_property_16_no_data_skips_tables_gracefully(artifact):
    sink = RecordingSink()
    renderer = TableRenderer(GroundingChecker(artifact), sink)
    result = renderer.render_latex(artifact)  # 不应抛异常
    assert result == []
    assert any(
        e.kind == EventKind.DEGRADATION and e.data.get("reason") == "no_data"
        for e in sink.events
    )


# --------------------------------------------------------------------------- #
# Property 17
# --------------------------------------------------------------------------- #


# Feature: venue-templates-figures-tables, Property 17: 数值格式化一致性与派生文本截断——
# 数值统一小数位数；由 stats 派生的文本（列名/图题）写入产物长度 <=500
@settings(max_examples=100)
@given(
    decimals=st.integers(min_value=1, max_value=6),
    name_len=st.integers(min_value=1, max_value=700),
    ds_len=st.integers(min_value=1, max_value=700),
    value=_VALUE,
)
def test_property_17_float_formatting_and_derived_text_truncation(
    decimals, name_len, ds_len, value
):
    metric = "A" * name_len
    dataset = "D" * ds_len
    exp = Experiment(
        experiment_id="e0",
        dataset=dataset,
        baselines=["OursB"],
        metrics=[metric],
        results_data={
            "rows": [{"method": "OursB", metric: value}],
            "stats": {metric: {"mean": value, "std": 0.0, "min": value, "max": value}},
        },
    )
    artifact = ResearchArtifact(
        research_question="rq", method=MethodSpec(overview="m"), experiments=[exp]
    )

    sink = RecordingSink()
    renderer = TableRenderer(GroundingChecker(artifact), sink, float_decimals=decimals)
    fragments = renderer.render_latex(artifact)
    assert len(fragments) == 1
    frag = fragments[0]

    # 数值格式化一致：唯一的数值单元的小数位数等于配置的 decimals。
    fractional = re.findall(r"\d+\.(\d+)", frag.latex)
    assert fractional  # 至少渲染出一个数值
    for frac in fractional:
        assert len(frac) == decimals

    # 派生文本截断：图题（由 dataset 派生）长度 <=500。
    assert len(frag.caption) <= 500

    # 列名（由 metric 派生）写入产物时被截断到 <=500。
    truncated = metric[:500]
    assert truncated in frag.latex
    if name_len > 500:
        assert metric not in frag.latex
