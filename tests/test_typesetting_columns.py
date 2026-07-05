"""分栏排版原语测试（columns 作为可组合排版原语）。

覆盖：
- apply_columns 原语把 docx 设为 N 栏（单一实现，convert 与就地路径共用）；
- apply_typesetting 在 spec.columns 指定时施加分栏 + 段落级排版；
- set_typesetting 工具记录 columns 规格；
- 「给一个 docx，就地设成双栏」保结构路径（set_typesetting + polish_docx_inplace）。

python-docx 缺失时跳过。
"""

from __future__ import annotations

import copy

import pytest

from paper_agent.agent_platform.guardrail_gate import GuardrailGate
from paper_agent.agent_platform.models import AgentSession, Typesetting, WritingTask
from paper_agent.agent_platform.tools.context import ToolContext
from paper_agent.agent_platform.tools.docx_inplace_tool import (
    register_polish_docx_inplace,
)
from paper_agent.agent_platform.tools.typesetting_tool import register_set_typesetting
from paper_agent.elicitation import AutoElicitor
from paper_agent.providers.llm.base import LLMResponse
from paper_agent.tools.registry import ToolRegistry
from paper_agent.workspace.models import InputMode, PaperWorkspace
from paper_agent.workspace.repository import WorkspaceRepository


class _MemStore:
    def __init__(self):
        self._data = {}

    def load(self, wid):
        raw = self._data.get(wid)
        return PaperWorkspace.from_dict(raw) if raw else None

    def save(self, ws):
        self._data[ws.workspace_id] = copy.deepcopy(ws.to_dict())


class _LLM:
    def complete(self, messages, **opts):
        return LLMResponse(content="（不应被调用）")


def _ctx(tmp_path, ws=None):
    ws = ws or PaperWorkspace(workspace_id="w1", input_mode=InputMode.DRAFT_REVISION)
    repo = WorkspaceRepository(_MemStore())
    repo.create(ws)
    session = AgentSession(session_id="w1", workspace=ws, task=WritingTask("t"))
    return ToolContext(
        session=session, repo=repo, gate=GuardrailGate(),
        elicitor=AutoElicitor(), output_dir=str(tmp_path),
    )


def _cols_num(docx_path):
    docx = pytest.importorskip("docx")
    from docx.oxml.ns import qn

    document = docx.Document(str(docx_path))
    cols = document.sections[0]._sectPr.find(qn("w:cols"))
    return cols.get(qn("w:num")) if cols is not None else None


def test_apply_columns_primitive_sets_two_columns(tmp_path):
    docx = pytest.importorskip("docx")
    from paper_agent.export.typesetting import apply_columns

    path = tmp_path / "d.docx"
    document = docx.Document()
    document.add_paragraph("正文。")
    document.save(str(path))

    affected = apply_columns(str(path), 2)
    assert affected >= 1
    assert _cols_num(path) == "2"


def test_apply_typesetting_applies_columns_and_alignment(tmp_path):
    docx = pytest.importorskip("docx")
    from docx.enum.text import WD_ALIGN_PARAGRAPH

    from paper_agent.export.typesetting import apply_typesetting

    path = tmp_path / "d.docx"
    document = docx.Document()
    document.add_paragraph("一段足够长的正文用于验证排版应用。")
    document.save(str(path))

    apply_typesetting(str(path), Typesetting(alignment="justify", columns=2))
    assert _cols_num(path) == "2"
    reopened = docx.Document(str(path))
    assert reopened.paragraphs[0].alignment == WD_ALIGN_PARAGRAPH.JUSTIFY


def test_set_typesetting_records_columns(tmp_path):
    ctx = _ctx(tmp_path)
    registry = ToolRegistry()
    register_set_typesetting(registry, ctx)
    out = registry.call("set_typesetting", columns=2)
    assert "双栏" in out
    assert ctx.workspace.profile["typesetting"]["columns"] == 2


def test_docx_inplace_two_column_preserves_source(tmp_path):
    """给一个 docx，就地设成双栏（保结构）：产物双栏、原稿字节不变。"""
    docx = pytest.importorskip("docx")
    src = tmp_path / "orig.docx"
    document = docx.Document()
    document.add_paragraph("一段正文内容。")
    document.save(str(src))
    original = src.read_bytes()

    ws = PaperWorkspace(workspace_id="w1", input_mode=InputMode.DRAFT_REVISION)
    ws.profile["typesetting"] = {"columns": 2, "alignment": "justify"}
    ctx = _ctx(tmp_path, ws)
    registry = ToolRegistry()
    register_polish_docx_inplace(registry, ctx, _LLM())

    out = registry.call(
        "polish_docx_inplace", path=str(src), polish_language=False
    )
    assert "排版" in out
    produced = tmp_path / "orig_inplace.docx"
    assert _cols_num(produced) == "2"  # 就地设成双栏
    assert src.read_bytes() == original  # 原稿逐字节不变
