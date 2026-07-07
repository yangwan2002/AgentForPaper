"""float_figure_top：内联图 → 浮动锚定（页顶/跨栏满宽/上下环绕）的 XML 结构校验。"""

from __future__ import annotations

import pytest

docx = pytest.importorskip("docx")
pytest.importorskip("PIL")

from docx.oxml.ns import qn  # noqa: E402
from docx.shared import Inches  # noqa: E402
from PIL import Image  # noqa: E402

from paper_agent.export.docx_float import float_figure_top  # noqa: E402


def _doc_with_images(path, n=1, img_path=None):
    Image.new("RGB", (8, 6), (180, 120, 60)).save(img_path)
    d = docx.Document()
    for i in range(n):
        d.add_paragraph(f"正文段 {i}")
        d.add_paragraph().add_run().add_picture(img_path, width=Inches(1))
    d.save(str(path))


def _first_drawing_child(d):
    for para in d.paragraphs:
        for run in para.runs:
            dr = run._r.find(qn("w:drawing"))
            if dr is not None:
                return dr
    return None


def test_inline_becomes_anchor_top_page_fullwidth(tmp_path):
    img = tmp_path / "f.png"
    path = tmp_path / "a.docx"
    _doc_with_images(path, n=1, img_path=str(img))

    ok, msg = float_figure_top(str(path), index=1, span_columns=True)
    assert ok, msg

    d = docx.Document(str(path))
    dr = _first_drawing_child(d)
    assert dr is not None
    anchor = dr.find(qn("wp:anchor"))
    assert anchor is not None                          # inline 已变 anchor
    assert dr.find(qn("wp:inline")) is None
    # 垂直：相对页面、顶端对齐。
    pv = anchor.find(qn("wp:positionV"))
    assert pv.get("relativeFrom") == "page"
    assert pv.find(qn("wp:align")).text == "top"
    # 水平：居中。
    ph = anchor.find(qn("wp:positionH"))
    assert ph.find(qn("wp:align")).text == "center"
    # 上下型环绕。
    assert anchor.find(qn("wp:wrapTopAndBottom")) is not None
    # 满宽：extent cx 接近版心宽（EMU 很大）。
    cx = int(anchor.find(qn("wp:extent")).get("cx"))
    assert cx > 4_000_000                              # >~4.4cm，明显被放大到整宽


def test_no_span_keeps_size(tmp_path):
    img = tmp_path / "f.png"; path = tmp_path / "b.docx"
    _doc_with_images(path, n=1, img_path=str(img))
    d0 = docx.Document(str(path))
    orig_cx = int(_first_drawing_child(d0).find(qn("wp:inline")).find(qn("wp:extent")).get("cx"))

    ok, _ = float_figure_top(str(path), index=1, span_columns=False)
    assert ok
    d = docx.Document(str(path))
    cx = int(_first_drawing_child(d).find(qn("wp:anchor")).find(qn("wp:extent")).get("cx"))
    assert cx == orig_cx                               # 不满宽 → 尺寸不变


def test_force_next_page_sets_page_break_before(tmp_path):
    img = tmp_path / "f.png"; path = tmp_path / "c.docx"
    _doc_with_images(path, n=1, img_path=str(img))
    ok, _ = float_figure_top(str(path), index=1, force_next_page=True)
    assert ok
    d = docx.Document(str(path))
    # 找到含图的段落，应带 pageBreakBefore。
    found = False
    for para in d.paragraphs:
        for run in para.runs:
            if run._r.find(qn("w:drawing")) is not None:
                p_pr = para._p.find(qn("w:pPr"))
                assert p_pr is not None and p_pr.find(qn("w:pageBreakBefore")) is not None
                found = True
    assert found


def test_preserves_drawing_and_paragraph_counts(tmp_path):
    img = tmp_path / "f.png"; path = tmp_path / "d.docx"
    _doc_with_images(path, n=2, img_path=str(img))
    d0 = docx.Document(str(path))
    n_para0 = len(d0.paragraphs)
    n_draw0 = sum(1 for p in d0.paragraphs for r in p.runs if r._r.find(qn("w:drawing")) is not None)

    ok, _ = float_figure_top(str(path), index=2, span_columns=True)
    assert ok
    d = docx.Document(str(path))
    n_para = len(d.paragraphs)
    n_draw = sum(1 for p in d.paragraphs for r in p.runs if r._r.find(qn("w:drawing")) is not None)
    assert n_para == n_para0                            # 段落数不变（Preservation 安全）
    assert n_draw == n_draw0                            # 图形数不变


def test_bad_index_returns_false(tmp_path):
    img = tmp_path / "f.png"; path = tmp_path / "e.docx"
    _doc_with_images(path, n=1, img_path=str(img))
    ok, msg = float_figure_top(str(path), index=5)
    assert ok is False
    assert "未找到" in msg


def test_missing_file_returns_false(tmp_path):
    ok, msg = float_figure_top(str(tmp_path / "nope.docx"), index=1)
    assert ok is False


def test_second_image_selected(tmp_path):
    img = tmp_path / "f.png"; path = tmp_path / "g.docx"
    _doc_with_images(path, n=3, img_path=str(img))
    ok, _ = float_figure_top(str(path), index=2, span_columns=True)
    assert ok
    d = docx.Document(str(path))
    # 恰好第 2 张变 anchor，另外两张仍是 inline。
    anchors = inlines = 0
    for p in d.paragraphs:
        for r in p.runs:
            dr = r._r.find(qn("w:drawing"))
            if dr is None:
                continue
            anchors += 1 if dr.find(qn("wp:anchor")) is not None else 0
            inlines += 1 if dr.find(qn("wp:inline")) is not None else 0
    assert anchors == 1 and inlines == 2
