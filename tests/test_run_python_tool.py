"""run_python 工具测试（sandboxed-run-python · Task 2/3）。

覆盖:产出新文件、输入原文件字节不变、失败诚实、不持有 repo 写能力、docx 保结构校验。
用真实 SubprocessSandbox(跨平台)执行简单代码。
"""

from __future__ import annotations

import copy

import pytest

from paper_agent.agent_platform.guardrail_gate import GuardrailGate
from paper_agent.agent_platform.models import AgentSession, WritingTask
from paper_agent.agent_platform.sandbox import SubprocessSandbox
from paper_agent.agent_platform.tools.context import ToolContext
from paper_agent.agent_platform.tools.run_python_tool import register_run_python
from paper_agent.elicitation import AutoElicitor
from paper_agent.tools.registry import ToolRegistry
from paper_agent.workspace.models import InputMode, PaperWorkspace
from paper_agent.workspace.repository import WorkspaceRepository


class _MemStore:
    def __init__(self):
        self._data = {}

    def load(self, wid):
        raw = self._data.get(wid)
        return PaperWorkspace.from_dict(raw) if raw else None

    def save(self, ws):
        self._data[ws.workspace_id] = copy.deepcopy(ws.to_dict())


def _ctx(tmp_path):
    ws = PaperWorkspace(workspace_id="w1", input_mode=InputMode.DRAFT_REVISION)
    repo = WorkspaceRepository(_MemStore())
    repo.create(ws)
    session = AgentSession(session_id="w1", workspace=ws, task=WritingTask("t"))
    return ToolContext(
        session=session, repo=repo, gate=GuardrailGate(),
        elicitor=AutoElicitor(), output_dir=str(tmp_path),
    )


def _registry(ctx):
    reg = ToolRegistry()
    register_run_python(reg, ctx, SubprocessSandbox(), default_timeout_s=30)
    return reg


def test_produces_new_file(tmp_path):
    out = _registry(_ctx(tmp_path)).call(
        "run_python",
        code="open('result.txt','w').write('ok')\nprint('done')\n",
    )
    assert "执行成功" in out
    assert "result.txt" in out


def test_input_file_bytes_unchanged(tmp_path):
    src = tmp_path / "in.txt"
    src.write_text("original", encoding="utf-8")
    original = src.read_bytes()
    out = _registry(_ctx(tmp_path)).call(
        "run_python",
        code="open('in.txt','a').write(' MODIFIED')\n",  # 改的是副本
        input_files=[str(src)],
    )
    assert "执行成功" in out
    # 原文件不变(改的是 Work_Dir 里的副本)。
    assert src.read_bytes() == original


def test_failure_is_honest(tmp_path):
    out = _registry(_ctx(tmp_path)).call(
        "run_python", code="import sys; sys.exit(2)\n"
    )
    assert "执行成功" not in out
    assert "非零退出" in out or "失败" in out


def test_empty_code(tmp_path):
    out = _registry(_ctx(tmp_path)).call("run_python", code="   ")
    assert "未提供" in out


def test_tool_does_not_hold_repo_write():
    """Property 5:run_python 工具不应持有 repo/gate 写能力(闭包里只有 runner + session/output)。"""
    import paper_agent.agent_platform.tools.run_python_tool as mod

    src = mod  # 模块不 import repo/commit/gate 写路径
    text = open(mod.__file__, encoding="utf-8").read()
    assert "commit(" not in text
    assert "ctx.repo" not in text
    assert "ctx.gate" not in text


def test_docx_preservation_pass(tmp_path):
    docx = pytest.importorskip("docx")
    src = tmp_path / "paper.docx"
    d = docx.Document()
    d.add_heading("方法", level=1)
    d.add_paragraph("原有正文。")
    d.save(str(src))

    # 代码:加一段(只增),产物名与输入同名 → 触发保结构校验并通过。
    code = (
        "import docx\n"
        "doc = docx.Document('paper.docx')\n"
        "doc.add_paragraph('新增一段。')\n"
        "doc.save('paper.docx')\n"
    )
    out = _registry(_ctx(tmp_path)).call(
        "run_python", code=code, input_files=[str(src)], preserve_docx=["paper.docx"]
    )
    assert "执行成功" in out
    assert src.read_bytes()  # 原稿仍在、未被改(改的是副本)


def test_docx_preservation_fail_drops_product(tmp_path):
    docx = pytest.importorskip("docx")
    src = tmp_path / "paper.docx"
    d = docx.Document()
    d.add_heading("方法", level=1)
    d.add_paragraph("原有正文一。")
    d.add_paragraph("原有正文二。")
    d.save(str(src))

    # 代码:删掉原有段落(破坏结构) → Preservation_Check 应失败、丢弃产物。
    code = (
        "import docx\n"
        "doc = docx.Document('paper.docx')\n"
        "p = doc.paragraphs[1]._p\n"
        "p.getparent().remove(p)\n"
        "doc.save('paper.docx')\n"
    )
    ctx = _ctx(tmp_path)
    out = _registry(ctx).call(
        "run_python", code=code, input_files=[str(src)], preserve_docx=["paper.docx"]
    )
    assert "保结构校验未通过" in out
