"""表格列宽修复（_fix_table_widths）——消除双栏窄栏里的逐字符压缩。

不依赖 pandoc：直接用 python-docx 造带固定列宽的表格，验证被改成自适应。
"""

from __future__ import annotations

import pytest

from paper_agent.agent_platform.tools.convert_tool import _fix_table_widths


def test_fix_table_widths_sets_autofit(tmp_path):
    docx = pytest.importorskip("docx")
    from docx.oxml.ns import qn
    from docx.shared import Inches

    document = docx.Document()
    table = document.add_table(rows=2, cols=4)
    # 造固定列宽（模拟 pandoc 产出的固定 tcW）。
    for row in table.rows:
        for cell in row.cells:
            cell.width = Inches(0.3)
    path = str(tmp_path / "t.docx")
    document.save(path)

    _fix_table_widths(path)

    reopened = docx.Document(path)
    tbl = reopened.tables[0]._tbl
    # tblLayout 被设为 autofit。
    layout = tbl.tblPr.find(qn("w:tblLayout"))
    assert layout is not None and layout.get(qn("w:type")) == "autofit"
    # 所有单元格宽度类型改为 auto。
    tc_ws = list(tbl.iter(qn("w:tcW")))
    assert tc_ws  # 确有单元格宽度节点
    assert all(w.get(qn("w:type")) == "auto" for w in tc_ws)


def test_fix_table_widths_no_tables_is_noop(tmp_path):
    docx = pytest.importorskip("docx")
    document = docx.Document()
    document.add_paragraph("无表格的文档。")
    path = str(tmp_path / "n.docx")
    document.save(path)
    # 无表格时不报错。
    _fix_table_widths(path)
    assert docx.Document(path).paragraphs[0].text == "无表格的文档。"
