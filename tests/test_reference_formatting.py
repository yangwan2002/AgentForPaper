"""参考文献学术排版测试：悬挂缩进 + 单倍行距，且不被正文排版覆盖。

- format_reference_paragraph 原语：设负首行缩进(悬挂) + 单倍行距；
- DocxExporter 导出的参考文献段落套用该格式、用受保护样式；
- 正文 apply_typesetting(对齐/行距/首行缩进) 不覆盖参考文献段落。

python-docx 缺失时跳过。
"""

from __future__ import annotations

import pytest

from paper_agent.agent_platform.models import Typesetting
from paper_agent.workspace.models import (
    InputMode,
    OutlineNode,
    OutputFormat,
    PaperWorkspace,
    ReferenceEntry,
    SectionDraft,
)


def test_format_reference_paragraph_sets_hanging_and_single(tmp_path):
    docx = pytest.importorskip("docx")
    from docx.enum.text import WD_LINE_SPACING

    from paper_agent.export.typesetting import format_reference_paragraph

    document = docx.Document()
    para = document.add_paragraph("1. Author (2024). Title. src:id")
    format_reference_paragraph(para)

    fmt = para.paragraph_format
    # 悬挂缩进：左缩进为正、首行缩进为等值负数。
    assert fmt.left_indent is not None and fmt.left_indent.pt > 0
    assert fmt.first_line_indent is not None and fmt.first_line_indent.pt < 0
    assert abs(fmt.left_indent.pt + fmt.first_line_indent.pt) < 0.5
    assert fmt.line_spacing_rule == WD_LINE_SPACING.SINGLE


def _ws_with_refs():
    ws = PaperWorkspace(
        workspace_id="wref", input_mode=InputMode.DRAFT_REVISION,
        output_format=OutputFormat.DOCX,
    )
    ws.outline = [OutlineNode(section_id="s1", title="引言", order=0)]
    ws.section_drafts = {
        "s1": SectionDraft(
            section_id="s1", title="引言",
            content="本文方法参考了已有工作 [1]，取得进展。",
        )
    }
    ws.verified_references = [
        ReferenceEntry(
            id="1", title="A Study", authors=["Smith"], year=2024,
            source_id="d1", source="openalex", verified=True,
        )
    ]
    return ws


def test_exported_references_have_hanging_indent(tmp_path):
    docx = pytest.importorskip("docx")
    from paper_agent.export.docx import DocxExporter

    ws = _ws_with_refs()
    result = DocxExporter().export(ws, str(tmp_path))
    assert result.files
    document = docx.Document(result.files[0])

    # 找到参考文献条目段落（以 "1. " 开头）。
    ref_paras = [p for p in document.paragraphs if p.text.strip().startswith("1. ")]
    assert ref_paras, "未找到参考文献条目段落"
    fmt = ref_paras[0].paragraph_format
    assert fmt.first_line_indent is not None and fmt.first_line_indent.pt < 0  # 悬挂


def test_body_typesetting_does_not_override_reference_hanging(tmp_path):
    """对导出 docx 施加正文排版(两端对齐/首行缩进)后，参考文献悬挂缩进仍保留。"""
    docx = pytest.importorskip("docx")
    from paper_agent.export.docx import DocxExporter
    from paper_agent.export.typesetting import apply_typesetting

    ws = _ws_with_refs()
    path = DocxExporter().export(ws, str(tmp_path)).files[0]

    # 正文排版：两端对齐 + 首行缩进 2 字符 + 行距 22 磅（模拟 set_typesetting）。
    apply_typesetting(path, Typesetting(alignment="justify", first_line_indent="2ch", line_spacing=22))

    document = docx.Document(path)
    ref_paras = [p for p in document.paragraphs if p.text.strip().startswith("1. ")]
    assert ref_paras
    fmt = ref_paras[0].paragraph_format
    # 参考文献仍是悬挂缩进（负首行缩进），未被正文的正首行缩进覆盖。
    assert fmt.first_line_indent is not None and fmt.first_line_indent.pt < 0
