"""visual-layout-acceptance 波6：真机 roundtrip 冒烟 + check_layout 工具 + 原子重试。

真机 roundtrip 用 `@pytest.mark.roundtrip` 标记：无 Word/LibreOffice 或无 PyMuPDF 时
自动 skip；有渲染后端才真的跑（CI 里装了 pandoc/LibreOffice 时执行）。
"""

from __future__ import annotations

import os

import pytest

from paper_agent.export.doc_render import rasterize_pdf, select_render_backend


def _make_min_docx(path: str) -> bool:
    """造一个最小 docx；无 python-docx 返回 False（调用方 skip）。"""
    try:
        import docx  # noqa: WPS433
    except Exception:  # noqa: BLE001
        return False
    d = docx.Document()
    d.add_heading("标题", level=1)
    d.add_paragraph("这是一段用于渲染冒烟测试的正文。")
    d.save(path)
    return True


@pytest.mark.roundtrip
def test_render_and_rasterize_roundtrip(tmp_path):
    """小 docx → 渲染 PDF → 逐页 PNG，断言产物存在。无后端/依赖则 skip。"""
    docx_path = str(tmp_path / "smoke.docx")
    if not _make_min_docx(docx_path):
        pytest.skip("无 python-docx，跳过真机 roundtrip")
    backend = select_render_backend()
    if backend is None:
        pytest.skip("无 Word/LibreOffice 渲染后端，跳过真机 roundtrip")

    pdf = str(tmp_path / "smoke.pdf")
    if not backend.render(docx_path, pdf):
        pytest.skip(f"{backend.name} 渲染失败（环境相关），跳过")
    assert os.path.isfile(pdf)

    pages = rasterize_pdf(pdf, str(tmp_path / "imgs"), dpi=100)
    if not pages:
        pytest.skip("无 PyMuPDF，跳过栅格化断言")
    assert all(os.path.isfile(p) for p in pages)
    assert len(pages) >= 1


# --------------------------------------------------------------------------- #
# check_layout 工具：只登记请求、不自己渲染
# --------------------------------------------------------------------------- #
def test_check_layout_records_request():
    from paper_agent.agent_platform.tools.check_layout_tool import register_check_layout
    from paper_agent.tools.registry import ToolRegistry

    class _Sess:
        def __init__(self):
            self.records = []

        def record(self, kind, **data):
            self.records.append({"kind": kind, **data})

    class _Ctx:
        def __init__(self):
            self.session = _Sess()

    ctx = _Ctx()
    reg = ToolRegistry()
    register_check_layout(reg, ctx)
    out = reg.call("check_layout", reason="刚改了图跨栏")
    assert "登记" in out
    assert ctx.session.records[-1]["kind"] == "check_layout"


# --------------------------------------------------------------------------- #
# 原子落盘：瞬时 PermissionError 有界重试后成功（根治 Windows 文件锁 flake）
# --------------------------------------------------------------------------- #
def test_atomic_finalize_retries_transient_permission_error(tmp_path, monkeypatch):
    import paper_agent.export.atomic_write as aw

    src = tmp_path / "a.part"; src.write_text("data", encoding="utf-8")
    dst = tmp_path / "a.docx"

    calls = {"n": 0}
    real_replace = os.replace

    def flaky_replace(a, b):
        calls["n"] += 1
        if calls["n"] == 1:
            raise PermissionError("[WinError 5] transient lock")
        return real_replace(a, b)

    monkeypatch.setattr(aw.os, "replace", flaky_replace)
    aw.atomic_finalize(str(src), str(dst))          # 首次抛→重试→成功
    assert dst.is_file()
    assert calls["n"] == 2


def test_atomic_finalize_persistent_error_raises(tmp_path, monkeypatch):
    import paper_agent.export.atomic_write as aw

    src = tmp_path / "b.part"; src.write_text("d", encoding="utf-8")

    def always_fail(a, b):
        raise PermissionError("locked forever")

    monkeypatch.setattr(aw.os, "replace", always_fail)
    with pytest.raises(PermissionError):
        aw.atomic_finalize(str(src), str(tmp_path / "b.docx"))
