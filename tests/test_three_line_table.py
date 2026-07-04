"""三线表样式（_apply_three_line_table_style）——学术论文标准表格边框。

不依赖 pandoc：用 python-docx 造表格、套三线表样式，验证边框设置。
"""

from __future__ import annotations

import pytest

from paper_agent.agent_platform.tools.convert_tool import (
    _apply_three_line_table_style,
)


def test_three_line_style_borders(tmp_path):
    docx = pytest.importorskip("docx")
    from docx.oxml.ns import qn

    document = docx.Document()
    table = document.add_table(rows=3, cols=4)
    path = str(tmp_path / "t.docx")
    document.save(path)

    _apply_three_line_table_style(path)

    reopened = docx.Document(path)
    tbl = reopened.tables[0]._tbl
    borders = tbl.tblPr.find(qn("w:tblBorders"))
    assert borders is not None

    def _val(name):
        el = borders.find(qn(f"w:{name}"))
        return el.get(qn("w:val")) if el is not None else None

    # 顶/底为实线，竖线与内部横线全为 none（三线表特征）。
    assert _val("top") == "single"
    assert _val("bottom") == "single"
    assert _val("left") == "none"
    assert _val("right") == "none"
    assert _val("insideV") == "none"
    assert _val("insideH") == "none"

    # 表头行（第一行）单元格底部有细线（\midrule）。
    first_row = reopened.tables[0].rows[0]
    for cell in first_row.cells:
        tc_borders = cell._tc.get_or_add_tcPr().find(qn("w:tcBorders"))
        assert tc_borders is not None
        bottom = tc_borders.find(qn("w:bottom"))
        assert bottom is not None and bottom.get(qn("w:val")) == "single"


def test_three_line_style_no_tables_noop(tmp_path):
    docx = pytest.importorskip("docx")
    document = docx.Document()
    document.add_paragraph("无表格。")
    path = str(tmp_path / "n.docx")
    document.save(path)
    _apply_three_line_table_style(path)  # 不报错
    assert docx.Document(path).paragraphs[0].text == "无表格。"
