"""docx 宽表适配：小号字体 + 紧凑边距 + 占满宽度 + 宽表跨双栏（补回 pandoc 丢失的
LaTeX 表格排版，消除多列表在窄栏里逐字符折行）。"""

from __future__ import annotations

import pytest

docx = pytest.importorskip("docx")
from docx.oxml.ns import qn  # noqa: E402

from paper_agent.agent_platform.tools.convert_tool import (  # noqa: E402
    _WIDE_TABLE_MIN_COLS,
    _compact_tables,
    _set_two_columns,
    _span_wide_tables,
)


def _make_doc(path, *table_cols):
    """造一个含若干指定列数表格的 docx（表间夹正文段），返回路径。"""
    d = docx.Document()
    for ncols in table_cols:
        t = d.add_table(rows=2, cols=ncols)
        for row in t.rows:
            for cell in row.cells:
                cell.text = "x"
        d.add_paragraph("正文段")
    d.save(str(path))
    return str(path)


def _para_sectpr_cols(path):
    """返回文档中段落级 sectPr 的 cols num 列表（顺序）。"""
    d = docx.Document(path)
    body = d.element.body
    xpath = qn("w:p") + "/" + qn("w:pPr") + "/" + qn("w:sectPr")
    out = []
    for sect in body.findall(xpath):
        cols = sect.find(qn("w:cols"))
        out.append(cols.get(qn("w:num")) if cols is not None else None)
    return out


def test_compact_tables_sets_small_font_margins_and_width(tmp_path):
    path = _make_doc(tmp_path / "a.docx", 8, 4)
    n = _compact_tables(path)
    assert n == 2
    d = docx.Document(path)
    for table in d.tables:
        # 字体缩到 9pt。
        run = table.rows[0].cells[0].paragraphs[0].runs[0]
        assert run.font.size.pt == 9.0
        tbl_pr = table._tbl.tblPr
        # 单元格左右边距收紧到 40 twips。
        mar = tbl_pr.find(qn("w:tblCellMar"))
        assert mar is not None
        assert mar.find(qn("w:left")).get(qn("w:w")) == "40"
        # 表宽为固定布局 dxa（按内容比例分列宽，取代 autofit 乱猜）。
        tbl_w = tbl_pr.find(qn("w:tblW"))
        assert tbl_w.get(qn("w:type")) == "dxa"
        assert int(tbl_w.get(qn("w:w"))) > 0
        assert tbl_pr.find(qn("w:tblLayout")).get(qn("w:type")) == "fixed"


def test_span_wide_only_wraps_tables_at_or_above_threshold(tmp_path):
    # 一张 8 列（宽，应跨栏）、一张 4 列（窄，不跨栏）。
    path = _make_doc(tmp_path / "b.docx", _WIDE_TABLE_MIN_COLS + 2, 4)
    _set_two_columns(path)
    spanned = _span_wide_tables(path)
    assert spanned == 1  # 只有宽表被跨栏
    # 宽表被裹进「前 2 栏 / 后 1 栏」的连续分节岛。
    assert _para_sectpr_cols(path) == ["2", "1"]


def test_narrow_table_not_spanned(tmp_path):
    path = _make_doc(tmp_path / "c.docx", 4, 3)
    _set_two_columns(path)
    spanned = _span_wide_tables(path)
    assert spanned == 0
    assert _para_sectpr_cols(path) == []  # 无新增段落级分节符


def test_content_proportional_widths_and_nowrap(tmp_path):
    """按内容比例分列宽（固定布局）+ 短单元格禁折行，取代 Word autofit 乱猜。"""
    from paper_agent.agent_platform.tools.convert_tool import _compact_tables

    d = docx.Document()
    t = d.add_table(rows=2, cols=3)
    # 第 2 列内容明显更长 → 应分到更宽的列宽。
    t.rows[0].cells[0].text = "N"
    t.rows[0].cells[1].text = "共视比例区间"
    t.rows[0].cells[2].text = "值"
    t.rows[1].cells[0].text = "174"
    t.rows[1].cells[1].text = "s in [0.010, 0.050]"
    t.rows[1].cells[2].text = "0.856"
    path = str(tmp_path / "w.docx")
    d.save(path)

    _compact_tables(path)

    d2 = docx.Document(path)
    tbl = d2.tables[0]._tbl
    # 固定布局。
    assert tbl.tblPr.find(qn("w:tblLayout")).get(qn("w:type")) == "fixed"
    grid = tbl.find(qn("w:tblGrid"))
    widths = [int(gc.get(qn("w:w"))) for gc in grid.findall(qn("w:gridCol"))]
    assert len(widths) == 3
    # 最长内容的第 2 列列宽 > 短列（N / 值）。
    assert widths[1] > widths[0] and widths[1] > widths[2]
    # 短无空格单元格（"N"/"174"/"值"/"0.856"）应带 noWrap。
    c0 = d2.tables[0].rows[1].cells[0]._tc
    assert c0.get_or_add_tcPr().find(qn("w:noWrap")) is not None


def test_section_text_width_twips_default(tmp_path):
    """取不到 section 尺寸时回退到合理默认，且正常文档返回版心宽区间内的值。"""
    from paper_agent.agent_platform.tools.convert_tool import _section_text_width_twips

    d = docx.Document()
    w = _section_text_width_twips(d)
    assert 1440 <= w <= 20000


def test_span_preserves_page_size_from_body_sectpr(tmp_path):
    path = _make_doc(tmp_path / "d.docx", 8)
    _set_two_columns(path)
    _span_wide_tables(path)
    d = docx.Document(path)
    body = d.element.body
    xpath = qn("w:p") + "/" + qn("w:pPr") + "/" + qn("w:sectPr")
    for sect in body.findall(xpath):
        # 新分节沿用了页面尺寸（不回落到默认纸张）。
        assert sect.find(qn("w:pgSz")) is not None
        # 连续分节，不产生分页。
        assert sect.find(qn("w:type")).get(qn("w:val")) == "continuous"
