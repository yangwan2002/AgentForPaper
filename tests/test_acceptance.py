"""确定性验收检查函数单测（Task 1）。

覆盖：乱码检测（真实 GBK↔latin1 样本 + 正常中文/英文对照）、排版核对、引用闭合
悬空/冗余、文献数量/年限边界。全部为无 LLM 的纯函数，结果确定可复现。
"""

from __future__ import annotations

import pytest

from paper_agent.agent_platform.acceptance import (
    AcceptanceReport,
    TaskRequirements,
    check_citation_closure,
    check_quantity,
    check_recency,
    check_typesetting_applied,
    cited_reference_ids,
    detect_mojibake,
)
from paper_agent.agent_platform.acceptance import AcceptanceFinding
from paper_agent.workspace.models import (
    InputMode,
    OutlineNode,
    PaperWorkspace,
    ReferenceEntry,
    SectionDraft,
)


# --------------------------------------------------------------------------- #
# 乱码检测
# --------------------------------------------------------------------------- #

def test_detect_mojibake_normal_chinese_is_clean():
    ok, _ = detect_mojibake("这是一段正常的中文论文摘要，包含标准的汉字与标点。")
    assert ok is False


def test_detect_mojibake_normal_english_is_clean():
    ok, _ = detect_mojibake("This is a perfectly normal English abstract with numbers 123.")
    assert ok is False


def test_detect_mojibake_utf8_read_as_latin1():
    # 中文经 UTF-8 编码后被误按 latin-1 解码 —— 真实乱码场景。
    garbled = "这是中文".encode("utf-8").decode("latin-1")
    is_bad, evidence = detect_mojibake(garbled)
    assert is_bad is True
    assert evidence


def test_detect_mojibake_replacement_char():
    is_bad, evidence = detect_mojibake("部分内容\ufffd\ufffd丢失")
    assert is_bad is True
    assert "FFFD" in evidence


def test_detect_mojibake_empty_is_clean():
    ok, _ = detect_mojibake("")
    assert ok is False


# --------------------------------------------------------------------------- #
# 引用闭合
# --------------------------------------------------------------------------- #

def _ws_with_refs() -> PaperWorkspace:
    ws = PaperWorkspace(workspace_id="w", input_mode=InputMode.GENERATION)
    ws.outline = [OutlineNode(section_id="s1", title="Intro", order=0)]
    ws.verified_references = [
        ReferenceEntry(id="1", title="A", authors=["X"], year=2020, source_id="d1", verified=True),
        ReferenceEntry(id="2", title="B", authors=["Y"], year=2021, source_id="d2", verified=True),
    ]
    return ws


def test_citation_closure_all_cited_passes():
    ws = _ws_with_refs()
    ws.section_drafts = {
        "s1": SectionDraft(section_id="s1", title="Intro", content="text [1] and [2]")
    }
    finding = check_citation_closure(ws)
    assert finding.ok is True


def test_citation_closure_redundant_still_passes_with_note():
    ws = _ws_with_refs()
    ws.section_drafts = {
        "s1": SectionDraft(section_id="s1", title="Intro", content="only cite [1]")
    }
    finding = check_citation_closure(ws)
    assert finding.ok is True
    assert "未被引用" in finding.detail


def test_citation_closure_dangling_fails_and_healable():
    ws = _ws_with_refs()
    ws.section_drafts = {
        "s1": SectionDraft(section_id="s1", title="Intro", content="cite [1] and [99]")
    }
    finding = check_citation_closure(ws)
    assert finding.ok is False
    assert finding.healable is True
    assert "99" in finding.detail


def test_cited_reference_ids_merges_text_and_recorded():
    ws = _ws_with_refs()
    ws.section_drafts = {
        "s1": SectionDraft(
            section_id="s1", title="Intro", content="text [1]",
            cited_reference_ids=["2"],
        )
    }
    assert cited_reference_ids(ws) == {"1", "2"}


# --------------------------------------------------------------------------- #
# 数量 / 年限
# --------------------------------------------------------------------------- #

def test_check_quantity_within_range():
    ws = _ws_with_refs()
    assert check_quantity(ws, lo=1, hi=5).ok is True


def test_check_quantity_below_min():
    ws = _ws_with_refs()
    f = check_quantity(ws, lo=3)
    assert f.ok is False and f.healable is True


def test_check_quantity_above_max():
    ws = _ws_with_refs()
    f = check_quantity(ws, hi=1)
    assert f.ok is False


def test_check_quantity_no_bounds_skips():
    ws = _ws_with_refs()
    assert check_quantity(ws).ok is True


def test_check_recency_all_recent():
    ws = _ws_with_refs()
    assert check_recency(ws, min_year=2019).ok is True


def test_check_recency_detects_stale():
    ws = _ws_with_refs()
    ws.verified_references.append(
        ReferenceEntry(id="3", title="Old", authors=["Z"], year=1999, source_id="d3", verified=True)
    )
    f = check_recency(ws, min_year=2000)
    assert f.ok is False
    assert "1999" in f.detail


def test_check_recency_no_min_skips():
    ws = _ws_with_refs()
    assert check_recency(ws).ok is True


# --------------------------------------------------------------------------- #
# 排版核对（依赖 python-docx，缺失则跳过）
# --------------------------------------------------------------------------- #

def test_check_typesetting_no_spec_passes():
    f = check_typesetting_applied("nonexistent.docx", {})
    assert f.ok is True


def test_check_typesetting_missing_file_reports_non_healable():
    pytest.importorskip("docx")
    f = check_typesetting_applied("does_not_exist_12345.docx", {"alignment": "justify"})
    assert f.ok is False
    assert f.healable is False


def test_check_typesetting_applied_matches(tmp_path):
    docx = pytest.importorskip("docx")
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.shared import Pt

    document = docx.Document()
    p = document.add_paragraph("正文段落内容，用于核对排版是否应用。")
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    p.paragraph_format.line_spacing = Pt(22)
    path = str(tmp_path / "typeset.docx")
    document.save(path)

    f = check_typesetting_applied(path, {"alignment": "justify", "line_spacing": 22})
    assert f.ok is True


def test_check_typesetting_detects_missing_alignment(tmp_path):
    docx = pytest.importorskip("docx")

    document = docx.Document()
    document.add_paragraph("未应用两端对齐的正文段落。")
    path = str(tmp_path / "plain.docx")
    document.save(path)

    f = check_typesetting_applied(path, {"alignment": "justify"})
    assert f.ok is False
    assert f.healable is True


# --------------------------------------------------------------------------- #
# 报告聚合
# --------------------------------------------------------------------------- #

def test_report_partitions_failures():
    report = AcceptanceReport(findings=[
        AcceptanceFinding(check="a", ok=True),
        AcceptanceFinding(check="b", ok=False, healable=True),
        AcceptanceFinding(check="c", ok=False, healable=False),
    ])
    assert report.passed is False
    assert [f.check for f in report.healable_failures] == ["b"]
    assert [f.check for f in report.blocking_failures] == ["c"]


def test_task_requirements_has_any():
    assert TaskRequirements(require_citation_closure=False).has_any() is False
    assert TaskRequirements(min_year=2020).has_any() is True
    assert TaskRequirements().has_any() is True  # 默认核对引用闭合


# --------------------------------------------------------------------------- #
# Task 3: AcceptanceChecker + AcceptanceLoop
# --------------------------------------------------------------------------- #

from dataclasses import dataclass

from paper_agent.agent_platform.acceptance import (
    AcceptanceChecker,
    AcceptanceLoop,
)


@dataclass
class _FakeSession:
    workspace: PaperWorkspace


def _ws_full() -> PaperWorkspace:
    ws = PaperWorkspace(workspace_id="w", input_mode=InputMode.GENERATION)
    ws.outline = [OutlineNode(section_id="s1", title="Intro", order=0)]
    ws.verified_references = [
        ReferenceEntry(id="1", title="A", authors=["X"], year=2020, source_id="d1", verified=True),
    ]
    ws.section_drafts = {
        "s1": SectionDraft(section_id="s1", title="Intro", content="正常中文正文，引用 [1]。")
    }
    return ws


def test_checker_all_pass_no_export_files():
    ws = _ws_full()
    report = AcceptanceChecker().check(ws, [], TaskRequirements())
    assert report.passed is True


def test_checker_dangling_citation_is_healable_failure():
    ws = _ws_full()
    ws.section_drafts["s1"].content = "引用 [1] 和 [99]"
    report = AcceptanceChecker().check(ws, [], TaskRequirements())
    assert report.passed is False
    assert report.healable_failures
    assert not report.blocking_failures


def test_checker_mojibake_in_export_is_blocking(tmp_path):
    ws = _ws_full()
    garbled = "这是中文".encode("utf-8").decode("latin-1")
    path = tmp_path / "out.md"
    path.write_text(garbled, encoding="utf-8")
    report = AcceptanceChecker().check(ws, [str(path)], TaskRequirements())
    assert report.passed is False
    assert report.blocking_failures  # 乱码不可自愈
    assert not any(f.healable for f in report.blocking_failures)


def test_checker_missing_expected_format(tmp_path):
    ws = _ws_full()
    md = tmp_path / "out.md"
    md.write_text("正常内容", encoding="utf-8")
    req = TaskRequirements(expected_format="docx")
    report = AcceptanceChecker().check(ws, [str(md)], req)
    assert any(f.check == "format" and not f.ok for f in report.findings)


# --- AcceptanceLoop ---

def test_loop_passes_immediately_when_clean():
    ws = _ws_full()
    session = _FakeSession(workspace=ws)
    loop = AcceptanceLoop(AcceptanceChecker(), export_fn=lambda w: [])
    outcome = loop.run(session, TaskRequirements(), max_heal_rounds=2)
    assert outcome.delivered is True
    assert outcome.heal_rounds == 0
    assert outcome.unresolved == []


def test_loop_heals_dangling_citation():
    ws = _ws_full()
    ws.section_drafts["s1"].content = "引用 [1] 和 [99]"
    session = _FakeSession(workspace=ws)

    def heal(sess, findings):
        # 模拟 Top_Agent 修正：删掉悬空标注（经单一写路径的行为在集成测试覆盖）。
        d = sess.workspace.section_drafts["s1"]
        d.content = d.content.replace(" 和 [99]", "")

    loop = AcceptanceLoop(AcceptanceChecker(), export_fn=lambda w: [], heal_fn=heal)
    outcome = loop.run(session, TaskRequirements(), max_heal_rounds=2)
    assert outcome.delivered is True
    assert outcome.heal_rounds == 1
    assert "citation_closure" in outcome.healed


def test_loop_bounded_when_heal_ineffective():
    ws = _ws_full()
    ws.section_drafts["s1"].content = "引用 [1] 和 [99]"
    session = _FakeSession(workspace=ws)
    calls = {"n": 0}

    def heal(sess, findings):
        calls["n"] += 1  # 无效修正：问题始终存在

    loop = AcceptanceLoop(AcceptanceChecker(), export_fn=lambda w: [], heal_fn=heal)
    outcome = loop.run(session, TaskRequirements(), max_heal_rounds=3)
    assert outcome.delivered is False
    assert outcome.heal_rounds == 3  # 有界终止
    assert calls["n"] == 3
    assert outcome.unresolved  # 诚实上报未解决项


def test_loop_does_not_heal_blocking_mojibake(tmp_path):
    ws = _ws_full()
    garbled = "这是中文".encode("utf-8").decode("latin-1")
    path = tmp_path / "out.md"
    path.write_text(garbled, encoding="utf-8")
    session = _FakeSession(workspace=ws)
    heal_calls = {"n": 0}

    def heal(sess, findings):
        heal_calls["n"] += 1

    loop = AcceptanceLoop(
        AcceptanceChecker(), export_fn=lambda w: [str(path)], heal_fn=heal
    )
    outcome = loop.run(session, TaskRequirements(), max_heal_rounds=2)
    assert outcome.delivered is False
    assert heal_calls["n"] == 0  # 乱码不可自愈 → 不触发 heal
    assert outcome.unresolved
